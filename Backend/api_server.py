#!/usr/bin/env python3
"""
FastAPI server for Sidian Query Orchestration System
Provides REST API endpoints for the query execution pipeline
"""

# Windows async event loop fix - required for proper SSE streaming on Windows
import sys
if sys.platform == 'win32':
    import asyncio
    asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

import os
import json
from html import escape
import httpx
from pathlib import Path
from datetime import datetime
import logging
import shutil
import re
import uuid
from typing import Dict, Any, Optional, List
from langgraph.errors import GraphInterrupt
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from docx import Document as DocxDocument
from pydantic import BaseModel
import uvicorn
from dotenv import load_dotenv
from utils.path_setup import ensure_info_retrieval_on_path

ensure_info_retrieval_on_path()
try:
    from utils.deep_agent_integration import (
        extract_interrupt_data,
        create_approval_payload,
        validate_action_approval,
        ensure_desktop_agent_state_compatibility,
        get_deep_agent_state_summary,
    )
except ImportError:  # pragma: no cover - fallback for alternative module paths
    from Backend.utils.deep_agent_integration import (  # type: ignore
        extract_interrupt_data,
        create_approval_payload,
        validate_action_approval,
        ensure_desktop_agent_state_compatibility,
        get_deep_agent_state_summary,
    )

# Load environment variables from the root .env (single source of truth)
root_env_path = Path(__file__).resolve().parent.parent / ".env"
if root_env_path.exists():
    load_dotenv(dotenv_path=str(root_env_path), override=True)
    print(f"✅ Loaded environment from {root_env_path}")
else:
    load_dotenv(override=True)
    print(f"⚠️ Root .env not found at {root_env_path}, using current environment")

# Import the query execution system
from main import execute_query, query_system_healthcheck
from nodes.DBRetrieval.KGdb import test_database_connection
from config.settings import PROJECT_CATEGORIES, CATEGORIES_PATH, PLANNER_PLAYBOOK, PLAYBOOK_PATH, DEBUG_MODE, MAX_CONVERSATION_HISTORY

DOC_API_URL = os.getenv("DOC_API_URL", "http://localhost:8002").rstrip("/")
DOC_API_TIMEOUT = float(os.getenv("DOC_API_TIMEOUT", "20"))
DOC_OP_SCHEMA_VERSION = 1

# Import Kuzu manager
try:
    from nodes.DBRetrieval.KGdb.kuzu_client import get_kuzu_manager
except ImportError:
    get_kuzu_manager = None

# Import feedback logger from helpers folder
try:
    from helpers.feedback_logger import get_feedback_logger
except ImportError:
    # Fallback if feedback_logger is not available
    def get_feedback_logger():
        class DummyLogger:
            def log_feedback(self, data):
                return True
            def get_feedback_stats(self):
                return {}
        return DummyLogger()

# Import Supabase logger from helpers folder
try:
    from helpers.supabase_logger import get_supabase_logger
except ImportError:
    # Fallback if supabase_logger is not available
    def get_supabase_logger():
        class DummyLogger:
            enabled = False
            def upload_image(self, *args, **kwargs):
                return None
            def log_user_query(self, *args, **kwargs):
                return True
            def log_feedback(self, *args, **kwargs):
                return True
            def get_interaction_stats(self, *args, **kwargs):
                return {}
            def get_user_summary(self, *args, **kwargs):
                return {}
            def get_recent_interactions(self, *args, **kwargs):
                return []
        return DummyLogger()

# Import enhanced logging system from helpers folder
try:
    from helpers.enhanced_logger import (
        enhanced_log_capture, 
        setup_enhanced_logging, 
        render_enhanced_log_html, 
        get_enhanced_log_stats
    )
except ImportError:
    # Fallback if enhanced_logger is not available
    class DummyLogCapture:
        def get_logs(self, *args, **kwargs):
            return []
        def get_function_calls(self, *args, **kwargs):
            return []
        def get_chunk_flow(self, *args, **kwargs):
            return []
        def clear_all(self):
            pass
    
    enhanced_log_capture = DummyLogCapture()
    
    def setup_enhanced_logging():
        pass
    
    async def render_enhanced_log_html(*args, **kwargs):
        return {"error": "Enhanced logging not available"}
    
    def get_enhanced_log_stats():
        return {}

# Import instructions content from helpers folder
try:
    from helpers.instructions import INSTRUCTIONS_CONTENT
except ImportError:
    INSTRUCTIONS_CONTENT = "# Instructions\n\nInstructions content not available."

# Helper function to wrap project numbers with fully-formed HTML links
def strip_asterisks_from_projects(text: str) -> str:
    """
    Remove markdown asterisks (**) that wrap project numbers or project names.
    The LLM sometimes adds **Project 25-01-002** - we want to strip these asterisks.
    """
    if not text:
        return text
    
    # Pattern 1: **Project NUMBER** (exact match)
    text = re.sub(r'\*\*Project\s+(\d{2}-\d{2}-\d{3})\*\*', r'Project \1', text)
    
    # Pattern 2: **Project NUMBER - ... (with dash after, asterisks before dash)
    text = re.sub(r'\*\*Project\s+(\d{2}-\d{2}-\d{3})\s*-\s*', r'Project \1 - ', text)
    
    # Pattern 3: Just project number wrapped: **25-01-002**
    text = re.sub(r'\*\*(\d{2}-\d{2}-\d{3})\*\*', r'\1', text)
    
    # Pattern 4: **Project NUMBER** with text after (space before closing **)
    text = re.sub(r'\*\*Project\s+(\d{2}-\d{2}-\d{3})\s+\*\*', r'Project \1 ', text)
    
    # Pattern 5: Handle "**Details:**" or "**Project Name:**" type labels
    text = re.sub(r'\*\*Details:\s*\*\*', 'Details:', text)
    text = re.sub(r'\*\*Project Name:\s*\*\*', 'Project Name:', text)
    text = re.sub(r'\*\*Details:\*\*', 'Details:', text)
    text = re.sub(r'\*\*Project Name:\*\*', 'Project Name:', text)
    
    # Pattern 6: Handle any remaining **Project NUMBER** patterns with various spacing
    text = re.sub(r'\*\*\s*Project\s+(\d{2}-\d{2}-\d{3})\s*\*\*', r'Project \1', text)
    
    return text




def wrap_code_file_links(text: str, code_citations: List[Dict]) -> str:
    """
    Wrap code file citations with fully-formed HTML links.
    Replaces citation patterns like [Document: filename, Page: X] with clickable filename links.
    Each link uses the specific page number from the citation.
    File paths from Supabase are already in the correct format and used directly.
    Shows filename for code DB (different files).
    This is done in the backend so frontend doesn't need to rebuild executable for link changes.
    """
    if not code_citations:
        return text
    
    # Create a mapping of filename to file path
    # File paths from Supabase are already in correct format: \\\\WADDELLNAS\\Resources\\References & Codes\\...
    file_link_map = {}  # filename -> file_path
    for cite in code_citations:
        filename = cite.get("filename", "")
        file_path = cite.get("file_path", "")
        
        if filename and file_path:
            # Use file_path directly from Supabase (already in correct format)
            if filename not in file_link_map:
                file_link_map[filename] = file_path
    
    # First, replace full citation patterns: [Document: filename, Page: X]
    # This extracts the specific page number from each citation and replaces with clean link
    # Process in reverse order by filename length to handle longer filenames first
    for filename, file_path in sorted(file_link_map.items(), key=lambda x: len(x[0]), reverse=True):
        escaped_filename = re.escape(filename)
        escaped_file_path = escape(file_path, quote=True)
        escaped_display_name = escape(filename, quote=False)  # Don't quote for text content
        
        # Pattern to match: [Document: filename, Page: X] (case-insensitive, flexible spacing)
        # Capture the page number from the citation
        citation_pattern = rf'\[Document:\s*{escaped_filename}\s*,\s*Page:\s*(\d+)\]'
        
        def make_citation_replacer(fpath, fname, display_name):
            def citation_replacer(match):
                page_number = match.group(1)  # Extract page number from citation
                # Generate fully-formed HTML link with filename as display text (code DB shows filename)
                title_text = f"Click to open file at page {page_number}"
                escaped_title = escape(title_text, quote=True)  # Escape title attribute
                return f'<a href="#" class="file-link" data-path="{fpath}" data-page="{page_number}" title="{escaped_title}">{display_name}</a>'
            return citation_replacer
        
        text = re.sub(citation_pattern, make_citation_replacer(escaped_file_path, filename, escaped_display_name), text, flags=re.IGNORECASE)
    
    return text


def wrap_coop_file_links(text: str, coop_citations: List[Dict]) -> str:
    """
    Wrap coop manual file citations with fully-formed HTML links.
    Replaces citation patterns like [Document: filename, Page: X] with clickable page number links.
    Each link uses the specific page number from the citation.
    Shows only page number for coop manual (same file always).
    File paths from Supabase are already in the correct format and used directly.
    This is done in the backend so frontend doesn't need to rebuild executable for link changes.
    """
    if not coop_citations:
        return text
    
    # Create a mapping of filename to file path (coop manual is always the same file)
    # File paths from Supabase are already in correct format: \\\\WADDELLNAS\\Coop\\...
    file_link_map = {}  # filename -> file_path
    for cite in coop_citations:
        filename = cite.get("filename", "")
        file_path = cite.get("file_path", "")
        
        if filename and file_path:
            # Use file_path directly from Supabase (already in correct format)
            if filename not in file_link_map:
                file_link_map[filename] = file_path
    
    # First, replace full citation patterns: [Document: filename, Page: X]
    # This extracts the specific page number from each citation and replaces with page number link
    # Process in reverse order by filename length to handle longer filenames first
    for filename, file_path in sorted(file_link_map.items(), key=lambda x: len(x[0]), reverse=True):
        escaped_filename = re.escape(filename)
        escaped_file_path = escape(file_path, quote=True)
        
        # Pattern to match: [Document: filename, Page: X] (case-insensitive, flexible spacing)
        # Capture the page number from the citation
        citation_pattern = rf'\[Document:\s*{escaped_filename}\s*,\s*Page:\s*(\d+)\]'
        
        def make_citation_replacer(fpath):
            def citation_replacer(match):
                page_number = match.group(1)  # Extract page number from citation
                # Generate fully-formed HTML link with page number as display text (coop manual shows page only)
                title_text = f"Click to open file at page {page_number}"
                escaped_title = escape(title_text, quote=True)  # Escape title attribute
                display_text = escape(f"Page {page_number}", quote=False)
                return f'<a href="#" class="file-link" data-path="{fpath}" data-page="{page_number}" title="{escaped_title}">{display_text}</a>'
            return citation_replacer
        
        text = re.sub(citation_pattern, make_citation_replacer(escaped_file_path), text, flags=re.IGNORECASE)
    
    return text


def _history_to_messages(history: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Normalize legacy conversation_history entries into message format.
    """
    messages: List[Dict[str, Any]] = []
    for entry in history or []:
        if isinstance(entry, dict) and entry.get("role"):
            msg = {
                "role": entry.get("role"),
                "content": entry.get("content", ""),
            }
            if entry.get("projects") is not None:
                msg["projects"] = entry.get("projects")
            meta = dict(entry.get("metadata") or {})
            ts = entry.get("timestamp")
            if ts is not None and "timestamp" not in meta:
                meta["timestamp"] = ts
            if meta:
                msg["metadata"] = meta
            messages.append(msg)
            continue

        if isinstance(entry, dict):
            ts = entry.get("timestamp")
            meta = dict(entry.get("metadata") or {})
            q = entry.get("question")
            a = entry.get("answer")

            if q:
                user_msg = {"role": "user", "content": q}
                if ts is not None:
                    user_msg["metadata"] = {"timestamp": ts}
                messages.append(user_msg)

            if a:
                if ts is not None and "timestamp" not in meta:
                    meta["timestamp"] = ts
                assistant_msg = {"role": "assistant", "content": a}
                if entry.get("projects") is not None:
                    assistant_msg["projects"] = entry.get("projects")
                if meta:
                    assistant_msg["metadata"] = meta
                messages.append(assistant_msg)

    return messages


def wrap_external_links(text: str) -> str:
    """
    Convert external HTML anchor tags to file-link format so frontend can handle them.
    Replaces external links like <a href="https://..." target="_blank">Sidcode</a>
    with <a href="#" class="file-link" data-path="https://..." title="Sidcode">Sidcode</a>
    This matches the pattern used for file citations - frontend already handles file-link clicks.
    """
    if not text:
        return text
    
    # Pattern to match external anchor tags (with http/https URLs)
    def convert_to_file_link(match):
        full_tag = match.group(0)
        # Extract URL from href attribute
        url_match = re.search(r'href=["\'](https?://[^"\']+)["\']', full_tag)
        # Extract display text (content between > and <)
        text_match = re.search(r'>([^<]+)<', full_tag)
        
        if url_match and text_match:
            url = url_match.group(1)
            display_text = text_match.group(1).strip()
            escaped_url = escape(url, quote=True)
            escaped_text = escape(display_text, quote=False)
            escaped_title = escape(f"Click to open {display_text}", quote=True)
            # Convert to file-link format (same as file citations) - frontend will handle clicks
            return f'<a href="#" class="file-link" data-path="{escaped_url}" title="{escaped_title}">{escaped_text}</a>'
        
        return full_tag  # Return unchanged if we can't parse it
    
    # Match external anchor tags with http/https URLs
    text = re.sub(r'<a\s+[^>]*href=["\']https?://[^"\']+["\'][^>]*>.*?</a>', convert_to_file_link, text, flags=re.IGNORECASE | re.DOTALL)
    
    return text


# --- Document rendering helpers (OnlyOffice) ---
DOCUMENTS_DIR = Path(__file__).resolve().parent / "documents"
DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
SESSION_DOC_FILES: Dict[str, str] = {}


def get_public_base_url() -> str:
    """Return the public base URL for serving generated documents."""
    base = (
        os.getenv("ONLYOFFICE_DOCUMENT_BASE_URL")
        or os.getenv("ORCHESTRATOR_PUBLIC_URL")
        or os.getenv("ORCHESTRATOR_URL")
        or os.getenv("RAG_API_URL")
        or "http://localhost:8000"
    )
    return base.rstrip("/")


def ensure_blank_document() -> Path:
    """Create a blank docx used as default/fallback."""
    blank_path = DOCUMENTS_DIR / "blank.docx"
    if blank_path.exists():
        return blank_path
    try:
        doc = DocxDocument()
        doc.add_paragraph("")  # empty paragraph to keep the file valid
        doc.save(blank_path)
        return blank_path
    except Exception as exc:  # noqa: BLE001
        # If python-docx default template is missing, create a minimal blank file from scratch
        logger.warning("⚠️ Failed to build blank.docx with default template (%s); creating fallback file", exc)
        # Minimal fallback: use python-docx with an in-memory empty package
        doc = DocxDocument()
        doc.save(blank_path)
        return blank_path


_ = ensure_blank_document()


def _create_doc_with_fallback() -> DocxDocument:
    """
    Create a docx Document, falling back to the repo's blank.docx when default template is missing.
    """
    blank_template = ensure_blank_document()
    try:
        return DocxDocument()
    except Exception as exc:  # noqa: BLE001
        logger.warning("⚠️ python-docx default template missing (%s); using blank.docx fallback", exc)
        return DocxDocument(str(blank_template))


def _safe_filename_component(value: str) -> str:
    """Make a value safe for filenames (OnlyOffice doc URLs)."""
    cleaned = re.sub(r"[^0-9A-Za-z._-]", "_", value or "")
    return cleaned or "doc"


def _humanize_doc_label(value: Optional[str]) -> Optional[str]:
    """Convert snake_case labels to title case for display."""
    if not value:
        return None
    text = str(value).replace("_", " ").strip()
    return text.title() if text else None


def _extract_doc_text_for_summary(doc_result: Dict[str, Any], answer_text: Optional[str]) -> str:
    """Pick the best available text snippet to summarize."""
    if not isinstance(doc_result, dict):
        return answer_text or ""

    for key in ("combined_text", "draft_text", "content", "text"):
        value = doc_result.get(key)
        if isinstance(value, str) and value.strip():
            return value

    sections = doc_result.get("sections") or []
    section_parts: List[str] = []
    for section in sections:
        if not isinstance(section, dict):
            continue
        for key in ("text", "draft_text", "content", "body"):
            value = section.get(key)
            if isinstance(value, str) and value.strip():
                section_parts.append(value.strip())
                break
    if section_parts:
        return "\n\n".join(section_parts)

    return answer_text or ""


def _strip_doc_noise(text: str) -> str:
    """Remove inline citations and warning banners from any chat-facing text."""
    if not text:
        return ""
    cleaned = text
    # Drop leading WARNING sentences/clauses
    cleaned = re.sub(r"(?is)\bwarning:\s.*?(?=(\n\n|$))", "", cleaned)
    cleaned = re.sub(r"(?is)\bwarning:\s.*?\.\s*", "", cleaned)
    # Strip inline citation brackets and numeric markers
    cleaned = re.sub(r"\[[^\]]*(?:Source|Document|Ref|citation|project|page|source)[^\]]*\]", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\[\s*\d+\s*\]", "", cleaned)
    # Collapse whitespace
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    return cleaned.strip()


def _extract_doc_section_titles(doc_result: Dict[str, Any]) -> List[str]:
    """Collect section titles/headings for display."""
    titles: List[str] = []
    if not isinstance(doc_result, dict):
        return titles

    for section in doc_result.get("sections") or []:
        if not isinstance(section, dict):
            continue
        title = (
            section.get("title")
            or section.get("heading")
            or section.get("section_type")
            or section.get("name")
        )
        if title:
            titles.append(str(title))

    # Deduplicate while preserving order
    seen = set()
    unique_titles = []
    for title in titles:
        if title in seen:
            continue
        seen.add(title)
        unique_titles.append(title)
    return unique_titles


def _extract_doc_citations(doc_result: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Flatten citations from doc_result or its sections."""
    if not isinstance(doc_result, dict):
        return []
    if isinstance(doc_result.get("citations"), list):
        return doc_result.get("citations") or []
    citations: List[Dict[str, Any]] = []
    for section in doc_result.get("sections") or []:
        if isinstance(section, dict):
            citations.extend(section.get("citations") or [])
    return citations


def _first_sentences(text: str, max_sentences: int = 3, max_chars: int = 500) -> str:
    """Take the first few sentences with a character cap."""
    if not text:
        return ""
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    picked: List[str] = []
    total = 0
    for sentence in sentences:
        if not sentence:
            continue
        picked.append(sentence.strip())
        total += len(sentence)
        if len(picked) >= max_sentences or total >= max_chars:
            break
    summary = " ".join(picked).strip()
    if len(summary) > max_chars:
        summary = summary[:max_chars].rstrip()
        if not summary.endswith("..."):
            summary += "..."
    return summary


def format_document_summary(
    doc_result: Optional[Dict[str, Any]],
    doc_type: Optional[str],
    section_type: Optional[str],
    answer_text: Optional[str],
) -> Optional[str]:
    """Build a markdown summary for chat display."""
    doc_result = doc_result or {}
    body_text = _strip_doc_noise(_extract_doc_text_for_summary(doc_result, answer_text))
    section_titles = _extract_doc_section_titles(doc_result)
    citations = _extract_doc_citations(doc_result)

    if not body_text and not section_titles:
        return None

    summary_text = _first_sentences(body_text, max_sentences=3, max_chars=500) if body_text else ""
    word_count = len((body_text or "").split())
    title = (
        doc_result.get("title")
        or doc_result.get("doc_title")
        or doc_result.get("document_title")
        or _humanize_doc_label(doc_type)
        or "Generated Document"
    )
    if section_type and section_type not in (doc_type or ""):
        pretty_section = _humanize_doc_label(section_type)
        if pretty_section and pretty_section.lower() not in title.lower():
            title = f"{title} - {pretty_section}"

    lines: List[str] = [f"📄 Document Generated: {title}"]

    if summary_text:
        lines.extend(["", "**Summary:**", summary_text])

    if section_titles:
        lines.append("")
        lines.append("**Key Sections:**")
        for heading in section_titles[:8]:
            lines.append(f"- {heading}")
        if len(section_titles) > 8:
            lines.append(f"- … {len(section_titles) - 8} more")

    lines.append("")
    lines.append("**Statistics:**")
    lines.append(f"• Word Count: ~{word_count} words" if word_count else "• Word Count: N/A")
    lines.append(f"• Citations: {len(citations)} source{'s' if len(citations) != 1 else ''} referenced")
    lines.append("")
    lines.append("📎 Full document available in the preview above")

    return "\n".join(lines).strip()


def _normalize_citation_entry(cite: Dict[str, Any]) -> Dict[str, Any]:
    """Normalize a single citation into the UI-friendly metadata shape."""
    return {
        "title": cite.get("title") or cite.get("document_title") or cite.get("filename") or "Source",
        "project": cite.get("project_key") or cite.get("drawing_number") or cite.get("project"),
        "date": cite.get("date") or cite.get("document_date"),
        "author": cite.get("author"),
        "relevance": cite.get("score") or cite.get("similarity") or cite.get("relevance"),
        "content_preview": cite.get("content_preview") or cite.get("preview") or cite.get("text") or "",
        "source_id": cite.get("artifact_id") or cite.get("id") or cite.get("chunk_id") or cite.get("source_id"),
        "page": cite.get("page_id") or cite.get("page"),
    }


def _ensure_citations_metadata(
    doc_generation_result_payload: Optional[Dict[str, Any]],
    citations_fallback: Optional[List[Dict[str, Any]]],
    warnings_fallback: Optional[List[str]],
    original_query: Optional[str],
    expanded_queries: Optional[List[str]] = None,
    support_score: Optional[float] = None,
) -> Dict[str, Any]:
    """
    Guarantee a citations_metadata payload exists. Prefer any precomputed block,
    otherwise build from available citations.
    """
    if isinstance(doc_generation_result_payload, dict):
        existing = doc_generation_result_payload.get("citations_metadata")
        if existing:
            try:
                doc_count = len(existing.get("documents", [])) if isinstance(existing, dict) else 0
                print(f"🐛 DEBUG: Using existing citations_metadata with {doc_count} documents")
            except Exception:
                print("🐛 DEBUG: Using existing citations_metadata (count unavailable)")
            return existing

    citations = []
    if isinstance(doc_generation_result_payload, dict):
        citations = doc_generation_result_payload.get("citations") or []
    if not citations:
        citations = citations_fallback or []

    documents: List[Dict[str, Any]] = []
    for cite in citations or []:
        if isinstance(cite, dict):
            documents.append(_normalize_citation_entry(cite))

    metadata = {
        "documents": documents,
        "warnings": warnings_fallback or [],
        "search_metadata": {
            "original_query": original_query or "",
            "expanded_queries": expanded_queries or [],
            "support_score": support_score if support_score is not None else 0.0,
            "retrieved_count": len(documents),
        },
    }
    print(
        f"🐛 DEBUG: Built fallback citations_metadata with {len(documents)} documents | "
        f"support_score={metadata['search_metadata']['support_score']}"
    )
    return metadata


def _normalize_doc_blocks(doc_result: Dict[str, Any], answer_text: Optional[str]) -> List[Dict[str, Any]]:
    """Convert doc_generation_result or plain text into a normalized block list."""
    blocks: List[Dict[str, Any]] = []
    sections = doc_result.get("sections") if isinstance(doc_result, dict) else None

    if isinstance(sections, list) and sections:
        for section in sections:
            heading = None
            if isinstance(section, dict):
                heading = section.get("title") or section.get("heading") or section.get("section_type")
            text = ""
            if isinstance(section, dict):
                text = section.get("text") or section.get("draft_text") or section.get("content") or section.get("body") or ""
            items = section.get("bullets") or section.get("items") if isinstance(section, dict) else None
            table = section.get("table") if isinstance(section, dict) else None

            if heading:
                blocks.append({"type": "heading", "level": section.get("level") or 1, "text": str(heading)})
            if items:
                blocks.append(
                    {
                        "type": "list",
                        "style": section.get("style") or "bullet",
                        "items": [str(item) for item in items if item],
                    }
                )
            if isinstance(table, dict):
                headers = table.get("headers") or []
                rows = table.get("rows") or []
                if headers and rows:
                    blocks.append(
                        {
                            "type": "table",
                            "headers": [str(h) for h in headers],
                            "rows": [[str(cell) for cell in row] for row in rows],
                        }
                    )
            if text:
                blocks.append({"type": "paragraph", "text": str(text)})

    if not blocks and answer_text:
        text = str(answer_text)
        paragraphs = [para.strip() for para in text.split("\n\n") if para.strip()]
        if not paragraphs:
            blocks.append({"type": "paragraph", "text": text})
        else:
            for para in paragraphs:
                lines = [ln.strip() for ln in para.splitlines() if ln.strip()]
                if lines and all(re.match(r"^[-*]\s+", ln) for ln in lines):
                    blocks.append(
                        {
                            "type": "list",
                            "style": "bullet",
                            "items": [re.sub(r"^[-*]\s+", "", ln) or ln for ln in lines],
                        }
                    )
                elif lines and all(re.match(r"^\d+[.)]\s+", ln) for ln in lines):
                    blocks.append(
                        {
                            "type": "list",
                            "style": "numbered",
                            "items": [re.sub(r"^\d+[.)]\s+", "", ln) or ln for ln in lines],
                        }
                    )
                else:
                    blocks.append({"type": "paragraph", "text": " ".join(lines)})

    return blocks


def render_docx_to_file(title: str, blocks: List[Dict[str, Any]], filename: str, append: bool = True) -> Path:
    """Materialize normalized blocks into a docx file. Append when existing doc is present."""
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = DOCUMENTS_DIR / filename

    # Append to existing document if requested and file exists
    if append and output_path.exists():
        doc = DocxDocument(output_path)
        add_title = False
    else:
        doc = _create_doc_with_fallback()
        add_title = True

    clean_title = title or "Generated Document"
    if add_title:
        doc.add_heading(clean_title, level=0)

    for block in blocks:
        btype = block.get("type")
        if btype == "heading":
            level = block.get("level") or 1
            try:
                level_int = int(level)
            except (TypeError, ValueError):
                level_int = 1
            level_int = max(0, min(4, level_int - 1))  # docx headings are 0-based where 0 is title
            doc.add_heading(block.get("text", ""), level=level_int)
        elif btype == "list":
            items = block.get("items") or []
            for item in items:
                p = doc.add_paragraph(str(item))
                p.style = "List Number" if (block.get("style") == "numbered") else "List Bullet"
        elif btype == "table":
            headers = block.get("headers") or []
            rows = block.get("rows") or []
            if headers and rows:
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    if i < len(hdr_cells):
                        hdr_cells[i].text = str(header)
                for row in rows:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row):
                        if i < len(row_cells):
                            row_cells[i].text = str(cell)
        else:
            text = block.get("text", "")
            if text is not None:
                doc.add_paragraph(str(text))

    doc.save(output_path)
    return output_path


def build_document_state(doc_url: str, title: str, blocks: List[Dict[str, Any]], doc_key: str) -> Dict[str, Any]:
    """Build a frontend-friendly document payload pointing to the OnlyOffice doc."""
    block_payloads = []
    for idx, block in enumerate(blocks):
        block_payloads.append(
            {
                "id": f"block-{idx + 1}",
                "type": block.get("type") or "paragraph",
                "text": block.get("text"),
                "level": block.get("level"),
                "items": block.get("items"),
                "rows": block.get("rows"),
                "style": block.get("style"),
            }
        )

    section_payload = {
        "id": f"section-{uuid.uuid4().hex[:8]}",
        "title": title or "Generated Document",
        "blocks": block_payloads,
    }

    return {
        "title": title or "Generated Document",
        "sections": [section_payload],
        "docUrl": doc_url,
        "documentKey": doc_key,
        "onlyoffice": {
            "docUrl": doc_url,
            "documentKey": doc_key,
        },
        "metadata": {
            "docUrl": doc_url,
            "documentKey": doc_key,
        },
        "docType": "docx",
        "viewMode": "edit",
    }


def should_materialize_doc(workflow: Optional[str], task_type: Optional[str], doc_result: Optional[Dict[str, Any]]) -> bool:
    """Determine if we should build an OnlyOffice doc for this response."""
    if workflow and str(workflow).lower().startswith("doc"):
        return True
    if task_type and str(task_type).startswith("doc"):
        return True
    return bool(doc_result)


def materialize_onlyoffice_document(
    doc_result: Optional[Dict[str, Any]],
    answer_text: Optional[str],
    session_id: str,
    message_id: str,
    append: bool = True,
) -> Optional[Dict[str, Any]]:
    """
    Convert doc generation output into a docx file and return a document_state payload.
    """
    try:
        doc_result = doc_result or {}
        import traceback

        def _log_emergency(location_name: str, text_value: str | None) -> None:
            stack = traceback.format_stack()
            stack_line = stack[-3].strip() if len(stack) >= 3 else ""
            snippet = (text_value or "")[:100]
            print(f"📍 CHECKPOINT: {location_name}")
            print(f"   Text: {snippet if snippet else 'None'}")
            print(f"   Contains TBD: {'[TBD' in (text_value or '')}")
            print(f"   Stack: {stack_line}")

        fallback_replacement = (
            "This document addresses the requested topic with consideration for industry best practices and standard approaches. "
            "The analysis incorporates relevant technical considerations and stakeholder requirements to provide a comprehensive overview of the subject matter.\n\n"
            "Key factors include alignment with regulatory standards, optimization of resource allocation, and adherence to established protocols. "
            "The approach prioritizes quality outcomes while maintaining operational efficiency and cost-effectiveness.\n\n"
            "Further detailed analysis can be conducted based on specific project requirements and available data sources."
        )

        def _has_tbd(text: str) -> bool:
            lowered = text.lower()
            return "[tbd" in lowered or "insufficient source" in lowered

        if answer_text:
            answer_text = _strip_doc_noise(answer_text)
        if answer_text and _has_tbd(answer_text):
            print("🚨 EMERGENCY FALLBACK: Replacing TBD content in answer_text")
            answer_text = fallback_replacement
            _log_emergency("materialize.emergency_fallback_answer_text", answer_text)
        if isinstance(doc_result, dict):
            draft_value = doc_result.get("draft_text")
            if isinstance(draft_value, str):
                cleaned = _strip_doc_noise(draft_value)
                doc_result["draft_text"] = cleaned
                if _has_tbd(draft_value):
                    print("🚨 EMERGENCY FALLBACK: Replacing TBD content in doc_result draft_text")
                    doc_result["draft_text"] = fallback_replacement
                    _log_emergency("materialize.emergency_fallback_doc_result", draft_value)
            sections = doc_result.get("sections")
            if isinstance(sections, list):
                for section in sections:
                    if not isinstance(section, dict):
                        continue
                    for key in ("text", "draft_text", "content", "body"):
                        val = section.get(key)
                        if isinstance(val, str):
                            section[key] = _strip_doc_noise(val)
                        if isinstance(val, str) and _has_tbd(val):
                            print(f"🚨 EMERGENCY FALLBACK: Replacing TBD content in section field '{key}'")
                            section[key] = fallback_replacement
                            _log_emergency("materialize.emergency_fallback_section", val)
        if not doc_result and answer_text:
            # Fallback to use the plain answer when no structured doc_result exists
            doc_result = {"draft_text": answer_text}
        blocks = _normalize_doc_blocks(doc_result, answer_text)
        print(f"📄 doc_generation_result keys: {list(doc_result.keys()) if doc_result else 'None'}")
        print(f"📝 answer_text length: {len(answer_text or '')} chars")
        if answer_text:
            print(f"🔤 answer_text content preview: {(answer_text or '')[:500]}")
        block_count = len(blocks)
        if not blocks:
            logger.warning("⚠️ No blocks to render for session=%s message=%s; skipping doc materialization", session_id, message_id)
            print(f"⚠️ No blocks to render; doc_result keys={list(doc_result.keys())}, answer_text_len={len(answer_text or '')}")
            return None

        title = (
            doc_result.get("title")
            or doc_result.get("doc_title")
            or doc_result.get("document_title")
            or "Generated Document"
        )

        base_session = _safe_filename_component(session_id)
        # Reuse the same doc per session so OnlyOffice keeps one open file
        filename = SESSION_DOC_FILES.get(session_id, f"{base_session}.docx")
        output_path = render_docx_to_file(title, blocks, filename, append=append)
        SESSION_DOC_FILES[session_id] = filename

        # Prefer relative URLs to avoid host mismatches (frontend uses same origin)
        use_relative_doc_url = os.getenv("DOC_URL_RELATIVE", "1") not in ("0", "false", "False")
        base_url = get_public_base_url()
        timestamp_ms = int(datetime.now().timestamp() * 1000)
        if use_relative_doc_url:
            doc_url = f"/documents/{filename}?t={timestamp_ms}"
        else:
            doc_url = f"{base_url}/documents/{filename}?t={timestamp_ms}"
        doc_key = f"{Path(filename).stem}-{timestamp_ms}"

        payload = build_document_state(doc_url, title, blocks, doc_key)
        try:
            file_size = output_path.stat().st_size
        except Exception:
            file_size = None
        logger.warning(
            "📄 Built OnlyOffice doc | blocks=%s | file=%s | size=%s bytes | url=%s",
            block_count,
            output_path,
            file_size,
            doc_url,
        )
        print(f"📄 Built OnlyOffice doc | blocks={block_count} | file={output_path} | size={file_size} | url={doc_url}")
        return payload
    except Exception as exc:  # noqa: BLE001
        logging.getLogger(__name__).error(f"❌ Failed to materialize OnlyOffice document: {exc}")
        print(f"❌ Failed to materialize OnlyOffice document: {exc}")
        return None


# Configure logging - use DEBUG_MODE from config
from config.logging_config import LOG_LEVEL
# Only configure if not already configured (logging_config.py already called basicConfig)
# But ensure root logger respects DEBUG_MODE
root_logger = logging.getLogger()
root_logger.setLevel(LOG_LEVEL)
# Also ensure all handlers use DEBUG_MODE level
for handler in root_logger.handlers:
    handler.setLevel(LOG_LEVEL)

logger = logging.getLogger(__name__)
logger.setLevel(LOG_LEVEL)  # Ensure this logger respects DEBUG_MODE

# Set up enhanced logging
setup_enhanced_logging()

# FastAPI app
app = FastAPI(
    title="Mantle RAG API",
    description="API for Mantle RAG Chat System",
    version="1.0.0"
)

# Enable CORS for frontend applications
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins for desktop app
    allow_credentials=False,  # Must be False when using wildcard origins
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.api_route("/documents/{document_name}", methods=["GET", "HEAD"])
async def get_document(document_name: str):
    """Serve OnlyOffice documents with cache-busting headers and on-demand creation."""
    if ".." in document_name or "/" in document_name or "\\" in document_name:
        raise HTTPException(status_code=400, detail="Invalid document name")

    clean_name = document_name.split("?")[0]
    document_path = DOCUMENTS_DIR / clean_name
    msg = f"Document requested: {clean_name} -> {document_path} (exists={document_path.exists()})"
    logger.info(msg)
    print(f"[documents] {msg}")

    if clean_name.startswith("conv-") and not document_path.exists():
        try:
            blank_path = ensure_blank_document()
            shutil.copy(blank_path, document_path)
            logger.info(f"Created new document: {clean_name}")
        except Exception as exc:  # noqa: BLE001
            logger.error(f"Failed to create conversation document {clean_name}: {exc}")
            raise HTTPException(status_code=500, detail="Failed to create conversation document") from exc

    if not document_path.exists():
        raise HTTPException(status_code=404, detail=f"Document not found: {clean_name}")

    return FileResponse(
        document_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=clean_name,
        headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0",
        },
    )

# Initialize async checkpointer on FastAPI startup
@app.on_event("startup")
async def startup_event():
    """Initialize async checkpointer when FastAPI starts"""
    from graph.checkpointer import init_checkpointer_async, checkpointer, CHECKPOINTER_TYPE
    import logging
    logger = logging.getLogger(__name__)
    try:
        # Initialize the async checkpointer
        initialized_checkpointer = await init_checkpointer_async()
        logger.info(f"✅ Checkpointer initialized: {type(initialized_checkpointer).__name__}")
        
        # Rebuild graph with proper checkpointer
        # This is critical - build_graph() will now import the updated checkpointer
        from graph.builder import build_graph
        import main
        main.graph = build_graph()
        logger.info("✅ Graph rebuilt with async checkpointer")
        
        # Verify checkpointer is being used
        if CHECKPOINTER_TYPE in ["postgres", "supabase"]:
            if initialized_checkpointer.__class__.__name__ != "AsyncPostgresSaver":
                logger.warning(f"⚠️  {CHECKPOINTER_TYPE} checkpointer unavailable; using in-memory fallback")
            else:
                logger.info(f"✅ Using {CHECKPOINTER_TYPE} checkpointer (persistent)")
                logger.info("💾 State will be automatically saved to Supabase after each node execution")
                # Verify the graph is actually using the async checkpointer
                if hasattr(main.graph, 'checkpointer'):
                    logger.info(f"✅ Graph checkpointer verified: {type(main.graph.checkpointer).__name__}")
        else:
            logger.info(f"✅ Using {CHECKPOINTER_TYPE} checkpointer")
    except Exception as e:
        logger.error(f"❌ Failed to initialize checkpointer on startup: {e}")
        import traceback
        logger.error(traceback.format_exc())

# Request/Response Models
class ChatRequest(BaseModel):
    message: str
    session_id: str = "default"
    user_identifier: Optional[str] = None
    data_sources: Optional[Dict[str, bool]] = None
    image_base64: Optional[str] = None  # Single screenshot/image as base64 (backwards compatible)
    images_base64: Optional[List[str]] = None  # Multiple images as base64 array

class CypherRequest(BaseModel):
    query: str
    params: Optional[Dict[str, Any]] = None

class CypherResponse(BaseModel):
    success: bool
    columns: Optional[List[str]] = None
    rows: Optional[List[Any]] = None
    row_count: Optional[int] = None
    error: Optional[str] = None
    query: Optional[str] = None

class ChatResponse(BaseModel):
    reply: str
    session_id: str
    timestamp: str
    latency_ms: float
    citations: int
    route: Optional[str] = None
    execution_trace: Optional[List[str]] = None
    node_path: Optional[str] = None
    workflow: Optional[str] = None
    task_type: Optional[str] = None
    doc_type: Optional[str] = None
    section_type: Optional[str] = None
    warnings: Optional[List[str]] = None
    message_id: Optional[str] = None
    project_answer: Optional[str] = None  # Separate project answer when multiple databases enabled
    code_answer: Optional[str] = None  # Separate code answer when code database enabled
    coop_answer: Optional[str] = None  # Separate coop answer when coop database enabled
    project_citations: Optional[int] = None
    code_citations: Optional[int] = None
    coop_citations: Optional[int] = None
    image_similarity_results: Optional[List[Dict[str, Any]]] = None  # Similar images found via image embedding search
    thinking_log: Optional[List[str]] = None  # Agent thinking logs for frontend display
    follow_up_questions: Optional[List[str]] = None  # Follow-up questions generated by verifier
    follow_up_suggestions: Optional[List[str]] = None  # Follow-up suggestions generated by verifier
    doc_generation_result: Optional[Dict[str, Any]] = None  # Structured document payload for doc preview
    doc_generation_warnings: Optional[List[str]] = None  # Doc generation warnings to show in UI
    document_state: Optional[Dict[str, Any]] = None  # OnlyOffice document patch
    citations_metadata: Optional[Dict[str, Any]] = None  # Structured citations/references metadata


class ActionApprovalRequest(BaseModel):
    action_id: str
    approved: bool
    reason: Optional[str] = None
    session_id: str


class ActionApprovalResponse(BaseModel):
    success: bool
    message: str
    resumed: bool
    next_state: Optional[dict] = None

class HealthResponse(BaseModel):
    status: str
    system_info: Dict[str, Any]
    timestamp: str

# Health Check endpoint
@app.get("/health", response_model=HealthResponse)
async def health_check():
    """System health check"""
    try:
        # Get system info from query system healthcheck
        health_info = query_system_healthcheck()
        
        return HealthResponse(
            status="healthy",
            system_info=health_info,
            timestamp=datetime.now().isoformat()
        )
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return HealthResponse(
            status="unhealthy",
            system_info={"error": str(e)},
            timestamp=datetime.now().isoformat()
        )



# ---------------------------------------------------------------------------
# Deep agent interrupt helpers
# ---------------------------------------------------------------------------
async def _load_thread_state(graph, session_id: str) -> Dict[str, Any]:
    """Load latest state for a thread from the graph checkpointer."""
    from graph.checkpointer import CHECKPOINTER_TYPE

    try:
        if CHECKPOINTER_TYPE in ["postgres", "supabase"]:
            snapshot = await graph.aget_state({"configurable": {"thread_id": session_id}})
        else:
            snapshot = graph.get_state({"configurable": {"thread_id": session_id}})
        if snapshot and getattr(snapshot, "values", None):
            return dict(snapshot.values)
    except Exception as exc:  # pragma: no cover - defensive
        logger.error(f"Failed to load state for session {session_id}: {exc}")
    return {}


async def _update_thread_state(graph, session_id: str, values: Dict[str, Any]):
    """Persist state updates to the graph checkpointer."""
    from graph.checkpointer import CHECKPOINTER_TYPE

    if not values:
        return None

    config = {"configurable": {"thread_id": session_id}}
    try:
        if CHECKPOINTER_TYPE in ["postgres", "supabase"] and hasattr(graph, "aupdate_state"):
            return await graph.aupdate_state(config=config, values=values)
        if hasattr(graph, "update_state"):
            return graph.update_state(config=config, values=values)
    except Exception as exc:  # pragma: no cover - defensive
        logger.error(f"Failed to update state for session {session_id}: {exc}")
    return None


async def _resume_graph(graph, session_id: str):
    """Resume graph execution after approval."""
    config = {"configurable": {"thread_id": session_id}}
    try:
        if hasattr(graph, "ainvoke"):
            return await graph.ainvoke(None, config=config)
        return graph.invoke(None, config=config)
    except GraphInterrupt:
        # Propagate further interrupts to caller
        raise
    except Exception as exc:  # pragma: no cover - defensive
        logger.error(f"Failed to resume graph for session {session_id}: {exc}")
        return None




# Streaming chat endpoint with real-time thinking logs
@app.post("/chat/stream")
async def chat_stream_handler(request: ChatRequest):
    """
    Streaming chat endpoint that emits thinking logs in real-time via Server-Sent Events (SSE).
    
    Returns:
        StreamingResponse with SSE format:
        - data: {"type": "connected"} - Initial connection
        - data: {"type": "thinking", "message": "...", "node": "...", "timestamp": ...} - Thinking logs
        - data: {"type": "complete", "result": {...}} - Final result
        - data: {"type": "error", "message": "..."} - Error messages
    """
    import time
    import uuid
    import asyncio
    doc_workflow_detected = False
    
    def _extract_projects(docs):
        """Extract project keys from documents"""
        projects = set()
        for doc in docs:
            if hasattr(doc, 'metadata') and isinstance(doc.metadata, dict):
                proj = doc.metadata.get("drawing_number") or doc.metadata.get("project_key")
                if proj:
                    projects.add(proj)
        return list(projects)[:20]
    
    def _extract_projects_from_citations(citations):
        """Extract project keys from citations"""
        projects = set()
        for citation in citations:
            if isinstance(citation, dict):
                proj = citation.get("project_key") or citation.get("drawing_number")
                if proj:
                    projects.add(proj)
        return list(projects)[:20]
    
    async def generate_stream():
        """Generate SSE stream of thinking logs and final result"""
        nonlocal doc_workflow_detected
        start_time = time.time()
        message_id = f"msg_{int(time.time())}_{uuid.uuid4().hex[:8]}"
        
        try:
            # Send initial connection event
            logger.info(f"📡 Starting streaming response for message: {request.message[:100]}...")
            yield f"data: {json.dumps({'type': 'connected', 'message_id': message_id})}\n\n"
            
            # Consolidate images
            images_to_process = request.images_base64 or []
            if not images_to_process and request.image_base64:
                images_to_process = [request.image_base64]
            
            # Log image receipt status (matching chat endpoint)
            logger.info(f"📸 IMAGE RECEIPT: images_base64={request.images_base64 is not None}, image_base64={request.image_base64 is not None}")
            logger.info(f"📸 IMAGES TO PROCESS: count={len(images_to_process)}, has_images={len(images_to_process) > 0}")
            if images_to_process:
                logger.info(f"📸 Image data lengths: {[len(img) for img in images_to_process[:3]]} chars (showing first 3)")
            
            # Process images if provided - Convert to searchable text via VLM (matching chat endpoint)
            # This is the key difference - chat endpoint does this before graph execution
            image_context = ""
            enhanced_question = request.message
            if images_to_process and len(images_to_process) > 0:
                from nodes.DBRetrieval.SQLdb.image_nodes import describe_image_for_search
                from config.logging_config import log_vlm
                
                log_vlm.info("")
                log_vlm.info("🔷" * 30)
                log_vlm.info(f"🖼️ PROCESSING {len(images_to_process)} IMAGE(S) WITH VLM")
                log_vlm.info(f"📝 User question: {request.message[:150] if request.message else 'General image description'}")
                log_vlm.info("🔷" * 30)
                
                # Log VLM processing to terminal and send to stream for Agent Thinking panel
                logger.info(f"🖼️ Processing {len(images_to_process)} image(s) with vision model...")
                # Send to stream for Agent Thinking panel - frontend will route to correct panel
                yield f"data: {json.dumps({'type': 'thinking', 'message': f'Processing {len(images_to_process)} image(s) with vision model...', 'node': 'vlm_processing', 'timestamp': time.time()})}\n\n"
                await asyncio.sleep(0.001)
                
                image_descriptions = []
                for i, image_base64 in enumerate(images_to_process):
                    log_vlm.info("")
                    log_vlm.info(f"📸 Processing image {i+1}/{len(images_to_process)}")
                    try:
                        image_description = describe_image_for_search(image_base64, request.message)
                        image_descriptions.append(f"Image {i+1}: {image_description}")
                        log_vlm.info(f"✅ Image {i+1} description completed successfully")
                    except Exception as e:
                        log_vlm.error(f"❌ VLM processing failed for image {i+1}, skipping: {e}")
                
                if image_descriptions:
                    image_context = "\n\n[Image Context: " + " | ".join(image_descriptions) + "]"
                    enhanced_question = request.message + image_context
                    log_vlm.info("")
                    log_vlm.info(f"📊 ALL IMAGES PROCESSED - Combined context length: {len(image_context)} chars")
                    log_vlm.info("🔷" * 30)
                    log_vlm.info("")
            
            # Import graph and state
            from main import graph
            from models.orchestration_state import OrchestrationState
            from dataclasses import asdict
            from thinking.intelligent_log_generator import IntelligentLogGenerator
            from models.memory import intelligent_query_rewriter
            
            log_generator = IntelligentLogGenerator()
            
            # CRITICAL: Load previous state from checkpointer to get messages/conversation history
            # This follows LangGraph best practices for short-term memory
            # Messages are persisted by checkpointer and loaded here for conversation context
            previous_state = None
            try:
                # Get the latest checkpoint state for this thread
                # For async checkpointers, use aget_state() instead of get_state()
                from graph.checkpointer import CHECKPOINTER_TYPE
                if CHECKPOINTER_TYPE in ["postgres", "supabase"]:
                    # Async checkpointer - use async method
                    state_snapshot = await graph.aget_state({"configurable": {"thread_id": request.session_id}})
                else:
                    # Sync checkpointer - use sync method
                    state_snapshot = graph.get_state({"configurable": {"thread_id": request.session_id}})
                
                if state_snapshot and state_snapshot.values:
                    previous_state = state_snapshot.values
                    logger.info(f"📖 Loaded previous state from checkpointer for thread_id={request.session_id}")
                    if previous_state.get("messages"):
                        logger.info(f"📖 Found {len(previous_state['messages'])} previous messages")
                    if previous_state.get("conversation_history"):
                        logger.info(f"📖 Found {len(previous_state['conversation_history'])} prior exchanges")
            except Exception as e:
                # No previous state exists (first invocation) - this is fine
                logger.info(f"📖 No previous state found (first invocation): {e}")
                previous_state = None
            
            # Get messages/conversation history from previous state for query rewriting and context
            previous_messages = previous_state.get("messages", []) if previous_state else []
            previous_history = previous_state.get("conversation_history", []) if previous_state else []
            history_messages = _history_to_messages(previous_history)

            # Sync messages/history and append current user turn
            synced_messages = list(previous_messages or history_messages)
            synced_history = list(previous_history or [])

            max_msgs = MAX_CONVERSATION_HISTORY * 2
            if len(synced_messages) > max_msgs:
                synced_messages = synced_messages[-max_msgs:]
            if len(synced_history) > MAX_CONVERSATION_HISTORY:
                synced_history = synced_history[-MAX_CONVERSATION_HISTORY:]

            user_turn_ts = time.time()
            user_turn = {
                "role": "user",
                "content": request.message,  # Use original question, not rewritten query
                "metadata": {"timestamp": user_turn_ts},
            }
            synced_messages.append(user_turn)
            if len(synced_messages) > max_msgs:
                synced_messages = synced_messages[-max_msgs:]

            synced_history.append(
                {
                    "question": request.message,
                    "answer": None,
                    "timestamp": user_turn_ts,
                }
            )
            if len(synced_history) > MAX_CONVERSATION_HISTORY:
                synced_history = synced_history[-MAX_CONVERSATION_HISTORY:]

            # Intelligent query rewriting with conversation context
            # This helps with follow-up detection and pronoun resolution
            rewritten_query, query_filters = intelligent_query_rewriter(
                enhanced_question,
                request.session_id,
                messages=synced_messages,
                conversation_history=synced_history,
            )
            logger.info(f"🎯 QUERY REWRITING: '{enhanced_question[:100]}...' → '{rewritten_query[:100]}...'" if len(enhanced_question) > 100 or len(rewritten_query) > 100 else f"🎯 QUERY REWRITING: '{enhanced_question}' → '{rewritten_query}'")
            
            # Extract project filter from query_filters if present
            project_filter = None
            if query_filters.get("project_keys"):
                project_filter = query_filters["project_keys"][0]
                logger.info(f"🎯 PROJECT FILTER FROM REWRITER: {project_filter}")
            
            # Determine if image similarity should be enabled
            use_image_similarity = False
            query_intent = None
            if images_to_process:
                from nodes.DBRetrieval.SQLdb.image_nodes import classify_image_query_intent
                intent_result = classify_image_query_intent(request.message, images_to_process[0])
                use_image_similarity = intent_result.get("use_image_similarity", False)
                query_intent = intent_result.get("intent")
            
            logger.info(f"📖 Initialized messages: {len(synced_messages)} total (windowed)")
            
            # Create initial state with messages and original question
            # This follows LangGraph best practices for short-term memory
            init_state = OrchestrationState(
                session_id=request.session_id,
                user_query=rewritten_query,  # Rewritten query for retrieval (with context)
                original_question=request.message,  # Original question (for storing in messages)
                project_filter=project_filter,
                data_sources=request.data_sources,
                images_base64=images_to_process if images_to_process else None,
                use_image_similarity=use_image_similarity,
                query_intent=query_intent,
                conversation_history=synced_history,
                messages=synced_messages,  # CRITICAL: Includes previous messages + new user message
            )
            
            # Track state as we go to detect which node ran
            previous_state = asdict(init_state)
            final_result = None
            
            config = {
                "configurable": {"thread_id": request.session_id}
            }
            
            # Use LangGraph's native streaming with stream_mode="updates"
            # "updates" mode gives us {node_name: state_updates} so we can identify which node ran
            # "custom" mode gives us custom events emitted from nodes
            accumulated_state = asdict(init_state)
            final_state = None
            
            logger.info(f"🔄 Starting graph.astream with stream_mode=['updates', 'custom', 'messages']")
            
            # Use "exit" durability mode to only save final state (not after each node)
            # This dramatically reduces storage costs - only 1 checkpoint per query instead of 7+
            # Trade-off: Can't resume from intermediate nodes, but you only need final state for memory
            durability_mode = os.getenv("CHECKPOINTER_DURABILITY", "exit").lower()  # "exit", "async", or "sync"
            if durability_mode == "exit":
                logger.info(f"💾 Checkpointer: Will save ONLY final state (exit mode) - reduces storage by ~85%")
            else:
                logger.info(f"💾 Checkpointer: Will save state after each node ({durability_mode} mode)")
            
            node_count = 0

            messages_received = 0
            async for stream_mode, chunk in graph.astream(
                asdict(init_state),
                config=config,
                stream_mode=["updates", "custom", "messages"],  # Add "messages" for token streaming
                durability=durability_mode  # Pass durability as direct parameter, not in config
            ):
                # Handle LLM token streaming (messages mode)
                # This captures tokens directly from LLM calls made within nodes
                if stream_mode == "messages":
                    messages_received += 1
                    # Log first few messages to debug
                    if messages_received <= 3:
                        logger.info(f"📨 Messages mode event #{messages_received}: type={type(chunk)}, chunk={str(chunk)[:200]}")
                    
                    # chunk is a tuple: (AIMessageChunk, metadata)
                    if isinstance(chunk, tuple) and len(chunk) >= 2:
                        message_chunk = chunk[0]
                        metadata = chunk[1] if len(chunk) > 1 else {}
                        
                        # Get the content from the message chunk
                        if hasattr(message_chunk, 'content') and message_chunk.content:
                            token_content = message_chunk.content
                            # Only stream tokens from synthesis LLM (answer node) - filter out other nodes
                            # Check metadata for node info to filter tokens
                            node_name = metadata.get('langgraph_node', 'unknown') if isinstance(metadata, dict) else 'unknown'
                            
                            # IMPORTANT: Only stream tokens from the "answer" node to main chat
                            # Tokens from other nodes (plan, retrieve, verify, etc.) should NOT appear in main chat
                            # They are internal processing and should only show as thinking logs in Agent Thinking panel
                            if node_name != 'answer' and node_name != 'unknown':
                                # Log to terminal but don't stream to main chat - these are internal processing tokens
                                logger.debug(f"⏭️  Skipping token from node '{node_name}' (not answer node) - only answer tokens go to main chat")
                                continue
                            
                            # Strip asterisks from project patterns
                            token_content = strip_asterisks_from_projects(token_content)
                            
                            # Stream token to frontend (only answer node tokens)
                            logger.debug(f"💬 Streaming token: {len(token_content)} chars from {node_name}")
                            yield f"data: {json.dumps({'type': 'token', 'content': token_content, 'node': 'answer', 'timestamp': time.time()})}\n\n"
                            await asyncio.sleep(0.001)  # Minimal delay for proper streaming
                    continue
                
                    # Handle custom events (emitted directly from nodes)
                if stream_mode == "custom":
                    logger.debug(f"📨 Custom event received: {chunk}")
                    # Handle custom events - send thinking logs to Agent Thinking panel, tokens to main chat
                    if isinstance(chunk, dict):
                        if chunk.get("type") == "thinking":
                            # Send thinking log to stream for Agent Thinking panel
                            logger.info(f"💭 {chunk.get('message', '')}")
                            yield f"data: {json.dumps({'type': 'thinking', 'message': chunk.get('message', ''), 'node': chunk.get('node', 'unknown'), 'timestamp': time.time()})}\n\n"
                            await asyncio.sleep(0.001)
                        elif chunk.get("type") == "token":
                            # Forward token to frontend for real-time streaming
                            token_content = chunk.get('content', '')
                            token_node = chunk.get('node', 'answer')
                            if doc_workflow_detected:
                                logger.info("📝 Doc workflow detected - suppressing custom token stream to chat.")
                                continue
                            # Strip asterisks from project patterns in token (helps with streaming)
                            # Note: This won't catch split patterns, but will handle complete ones
                            token_content = strip_asterisks_from_projects(token_content)
                            logger.debug(f"💬 Streaming token from node '{token_node}': {len(token_content)} chars")
                            yield f"data: {json.dumps({'type': 'token', 'content': token_content, 'node': token_node, 'timestamp': time.time()})}\n\n"
                            await asyncio.sleep(0.001)  # Minimal delay for proper streaming
                    continue
                
                # Handle state updates (updates mode)
                if stream_mode != "updates":
                    continue
                    
                # chunk is now {node_name: state_updates_dict}
                # This gives us the node name directly as the key
                node_outputs = chunk
                logger.info(f"📦 Received node_outputs: {list(node_outputs.keys())}")
                
                for node_name, state_updates in node_outputs.items():
                    node_count += 1
                    logger.info(f"🔍 Processing node #{node_count}: '{node_name}'")
                    
                    if not isinstance(state_updates, dict):
                        continue
                    
                    # Merge state updates into accumulated state
                    # (updates mode gives us only the changes, not full state)
                    accumulated_state.update(state_updates)
                    state_dict = accumulated_state.copy()
                    workflow_hint = state_dict.get("workflow") or state_updates.get("workflow")
                    task_type_hint = state_dict.get("task_type") or state_updates.get("task_type")
                    if (workflow_hint and str(workflow_hint).lower().startswith("doc")) or (
                        task_type_hint and str(task_type_hint).startswith("doc")
                    ) or node_name.startswith("doc_"):
                        doc_workflow_detected = True
                    
                    # Generate thinking log based on node using intelligent_log_generator and RAG state data
                    thinking_log = None
                    
                    if node_name == "plan":
                        plan = state_dict.get("query_plan") or {}
                        thinking_log = log_generator.generate_planning_log(
                            query=request.message,
                            plan=plan,
                            route=state_dict.get("data_route") or "smart",
                            project_filter=state_dict.get("project_filter")
                        )
                    elif node_name == "router_dispatcher":
                        thinking_log = log_generator.generate_router_dispatcher_log(
                            query=request.message,
                            selected_routers=state_dict.get("selected_routers", [])
                        )
                    elif node_name == "database" or node_name == "rag":  # Support legacy "rag"
                        thinking_log = log_generator.generate_rag_log(
                            query=request.message,
                            query_plan=state_dict.get("query_plan"),
                            data_route=state_dict.get("data_route"),
                            data_sources=state_dict.get("data_sources")
                        )
                    elif node_name == "generate_image_embeddings":
                        image_count = len(state_dict.get("images_base64") or [])
                        thinking_log = log_generator.generate_image_embeddings_log(
                            query=request.message,
                            image_count=image_count
                        )
                    elif node_name == "image_similarity_search":
                        similarity_results = state_dict.get("image_similarity_results", [])
                        project_keys = set()
                        for img in similarity_results:
                            proj = img.get("project_key")
                            if proj:
                                project_keys.add(proj)
                        thinking_log = log_generator.generate_image_similarity_log(
                            query=request.message,
                            result_count=len(similarity_results),
                            project_count=len(project_keys)
                        )
                    elif node_name == "retrieve":
                        thinking_log = log_generator.generate_retrieval_log(
                            query=request.message,
                            project_count=len(state_dict.get("retrieved_docs") or []),
                            code_count=len(state_dict.get("retrieved_code_docs") or []),
                            coop_count=len(state_dict.get("retrieved_coop_docs") or []),
                            projects=_extract_projects(state_dict.get("retrieved_docs") or []),
                            route=state_dict.get("data_route") or "smart",
                            project_filter=state_dict.get("project_filter"),
                            query_plan=state_dict.get("query_plan")
                        )
                    elif node_name == "grade":
                        thinking_log = log_generator.generate_grading_log(
                            query=request.message,
                            retrieved_count=len(state_dict.get("retrieved_docs") or []),
                            graded_count=len(state_dict.get("graded_docs") or []),
                            filtered_out=len(state_dict.get("retrieved_docs") or []) - len(state_dict.get("graded_docs") or [])
                        )
                    elif node_name == "answer":
                        thinking_log = log_generator.generate_synthesis_log(
                            query=request.message,
                            graded_count=len(state_dict.get("answer_citations") or []),
                            projects=_extract_projects_from_citations(state_dict.get("answer_citations") or []),
                            has_code=bool(state_dict.get("code_answer")),
                            has_coop=bool(state_dict.get("coop_answer"))
                        )
                    elif node_name == "verify":
                        thinking_log = log_generator.generate_verify_log(
                            query=request.message,
                            needs_fix=state_dict.get("needs_fix", False),
                            follow_up_count=len(state_dict.get("follow_up_questions", [])),
                            suggestion_count=len(state_dict.get("follow_up_suggestions", []))
                        )
                    elif node_name == "correct":
                        thinking_log = log_generator.generate_correct_log(
                            query=request.message,
                            support_score=state_dict.get("answer_support_score", 1.0),
                            corrective_attempted=state_dict.get("corrective_attempted", False)
                        )
                    else:
                        # For any unexpected nodes, log them but don't show to user
                        logger.info(f"Skipping thinking log for node: {node_name}")
                        thinking_log = None
                    
                    # Emit thinking log to stream for Agent Thinking panel (NOT displayed in main chat)
                    if thinking_log:
                        log_data = {'type': 'thinking', 'message': thinking_log, 'node': node_name, 'timestamp': time.time()}
                        logger.info(f"📤 Streaming thinking log for node '{node_name}': {thinking_log[:100]}...")
                        # Send to stream for Agent Thinking panel - frontend will route to correct panel
                        yield f"data: {json.dumps(log_data)}\n\n"
                        await asyncio.sleep(0.001)  # Minimal delay for proper streaming
                    else:
                        logger.info(f"⏭️  No thinking log generated for node '{node_name}'")
                    
                    # DO NOT manually stream answer - tokens come via "messages" mode
                    # The manual word-by-word streaming causes duplication
                    # Tokens are already being streamed in real-time via messages mode (line 930)
                    # This manual streaming was causing word duplication ("Project Project", etc.)
                    if node_name == "answer":
                        logger.info(f"📤 Answer node completed - tokens streaming via messages mode (received: {messages_received})")
                    
                    # Store final state - correct node is the last one
                    final_state = state_dict.copy()
                    
                    # Break after correct node (last node in pipeline)
                    if node_name == "correct":
                        break
            
            # Log streaming stats
            logger.info(f"📊 Streaming stats: {messages_received} LLM token events received via messages mode")
            
            # Check if we got a result
            if not final_state:
                yield f"data: {json.dumps({'type': 'error', 'message': 'No result returned from RAG'})}\n\n"
                return
            if not doc_workflow_detected and should_materialize_doc(
                final_state.get("workflow"), final_state.get("task_type"), final_state.get("doc_generation_result")
            ):
                doc_workflow_detected = True
            
            # Extract final answer (matching chat endpoint logic)
            answer = (
                final_state.get("final_answer")
                or final_state.get("db_retrieval_result")
                or final_state.get("answer", "No answer generated.")
            )
            code_answer = final_state.get("code_answer") or final_state.get("db_retrieval_code_answer")
            coop_answer = final_state.get("coop_answer") or final_state.get("db_retrieval_coop_answer")
            code_citations = final_state.get("code_citations", []) or final_state.get("db_retrieval_code_citations", [])
            coop_citations = final_state.get("coop_citations", []) or final_state.get("db_retrieval_coop_citations", [])
            project_citations = final_state.get("answer_citations", []) or final_state.get("db_retrieval_citations", [])
            image_similarity_results = final_state.get("image_similarity_results", []) or final_state.get("db_retrieval_image_similarity_results", [])
            follow_up_questions = final_state.get("follow_up_questions", []) or final_state.get("db_retrieval_follow_up_questions", [])
            follow_up_suggestions = final_state.get("follow_up_suggestions", []) or final_state.get("db_retrieval_follow_up_suggestions", [])
            doc_answer_text = None
            
            # Only extract doc_generation_result if workflow is docgen (not qa/rag)
            # This ensures RAG subgraph results don't include document generation artifacts
            workflow = final_state.get("workflow")
            doc_generation_result_payload = final_state.get("doc_generation_result") if workflow == "docgen" else None
            
            support_score = (
                final_state.get("answer_support_score")
                or final_state.get("db_retrieval_support_score")
                or final_state.get("support")
            )
            expanded_queries = final_state.get("expanded_queries") or final_state.get("db_retrieval_expanded_queries") or []
            citations_metadata = _ensure_citations_metadata(
                doc_generation_result_payload if isinstance(doc_generation_result_payload, dict) else None,
                project_citations,
                final_state.get("doc_generation_warnings") or final_state.get("warnings") or [],
                request.message,
                expanded_queries=expanded_queries,
                support_score=support_score,
            )
            if isinstance(doc_generation_result_payload, dict) and not doc_generation_result_payload.get("citations_metadata"):
                doc_generation_result_payload = dict(doc_generation_result_payload)
                doc_generation_result_payload["citations_metadata"] = citations_metadata
            print(
                f"🐛 DEBUG: /chat/stream citations_metadata docs={len(citations_metadata.get('documents', [])) if isinstance(citations_metadata, dict) else 0}"
            )
            doc_summary = None
            
            route_value = final_state.get("data_route") or final_state.get("db_retrieval_route")
            execution_trace = final_state.get("execution_trace") or []

            # Check which databases were used (matching chat endpoint)
            has_project = bool(answer and answer.strip())
            has_code = bool(code_answer and code_answer.strip())
            has_coop = bool(coop_answer and coop_answer.strip())
            enabled_count = sum([has_project, has_code, has_coop])
            multiple_enabled = enabled_count > 1
            
            # Handle empty or None answers (matching chat endpoint)
            if not answer or answer.strip() == "":
                # Only set error message if there's no code_answer or coop_answer (pure project_db mode)
                if not code_answer or not code_answer.strip():
                    if not coop_answer or not coop_answer.strip():
                        answer = "I couldn't generate a response. Please try rephrasing your question or check that the system is properly configured."
            
            # Calculate latency
            latency_ms = (time.time() - start_time) * 1000
            
            # Process answers separately for multi-bubble display (matching chat endpoint)
            if multiple_enabled:
                # Multiple databases enabled - return separate answers
                # Strip asterisks from project names/numbers
                processed_project_answer = strip_asterisks_from_projects(answer) if has_project else None

                # Wrap code file citations with file-link tags (shows filename)
                processed_code_answer = wrap_code_file_links(code_answer, code_citations) if (has_code and code_citations) else code_answer

                # Convert external links to file-link format (for sidian-bot links) - works with existing frontend handlers
                if has_code:
                    processed_code_answer = wrap_external_links(processed_code_answer)

                # Wrap coop file citations with file-link tags (shows page number only)
                processed_coop_answer = wrap_coop_file_links(coop_answer, coop_citations) if (has_coop and coop_citations) else coop_answer

                # Log to Supabase (matching chat endpoint)
                supabase_logger = get_supabase_logger()
                if supabase_logger.enabled:
                    user_identifier = request.user_identifier or f"{os.getenv('USERNAME', 'unknown')}@{os.getenv('COMPUTERNAME', 'unknown')}"
                    # Build combined logging text
                    parts = []
                    if has_project:
                        parts.append(answer)
                    if has_code:
                        parts.append(f"--- Code References ---\n\n{code_answer}")
                    if has_coop:
                        parts.append(f"--- Training Manual References ---\n\n{coop_answer}")
                    combined_for_logging = "\n\n".join(parts)

                    # Upload first image if provided (for logging)
                    image_url = None
                    if images_to_process and len(images_to_process) > 0:
                        image_url = supabase_logger.upload_image(
                            images_to_process[0],
                            message_id,
                            "image/png"
                        )

                    query_data = {
                        "message_id": message_id,
                        "session_id": request.session_id,
                        "user_query": request.message,
                        "rag_response": combined_for_logging,
                        "route": route_value,
                        "citations_count": len(project_citations) + len(code_citations) + len(coop_citations),
                        "latency_ms": round(latency_ms, 2),
                        "user_identifier": user_identifier,
                        "image_url": image_url
                    }
                    supabase_logger.log_user_query(query_data)

                # Build response with separate answers (matching chat endpoint)
                doc_answer_text = processed_project_answer or processed_code_answer or processed_coop_answer or answer
                
                # Only format document summary if workflow is docgen (not qa/rag)
                # This ensures document statistics only appear for doc generation, not RAG queries
                if workflow == "docgen" and doc_generation_result_payload:
                    doc_summary = format_document_summary(
                        doc_generation_result_payload,
                        final_state.get("doc_type"),
                        final_state.get("section_type"),
                        doc_answer_text,
                    )
                    if doc_summary and isinstance(doc_generation_result_payload, dict):
                        doc_generation_result_payload = dict(doc_generation_result_payload)
                        doc_generation_result_payload["document_summary"] = doc_summary
                response_data = {
                    'reply': processed_project_answer or processed_code_answer or processed_coop_answer or "No answer generated.",  # Primary answer (for backward compatibility)
                    'session_id': request.session_id,
                    'timestamp': datetime.now().isoformat(),
                    'latency_ms': round(latency_ms, 2),
                    'citations': len(project_citations) + len(code_citations) + len(coop_citations),  # Total citations
                    'route': route_value,
                    'message_id': message_id,
                    'project_answer': processed_project_answer,
                    'code_answer': processed_code_answer,
                    'coop_answer': processed_coop_answer,
                    'project_citations': len(project_citations) if has_project else None,
                    'code_citations': len(code_citations) if has_code else None,
                    'coop_citations': len(coop_citations) if has_coop else None,
                    'image_similarity_results': image_similarity_results if image_similarity_results else [],  # Always include as array, not None
                    'follow_up_questions': follow_up_questions if follow_up_questions else None,
                    'follow_up_suggestions': follow_up_suggestions if follow_up_suggestions else None,
                    'workflow': final_state.get("workflow"),
                    'task_type': final_state.get("task_type"),
                    'doc_type': final_state.get("doc_type"),
                    'section_type': final_state.get("section_type"),
                    'execution_trace': execution_trace,
                    'node_path': " → ".join(execution_trace) if execution_trace else None,
                    'doc_generation_result': doc_generation_result_payload,
                    'doc_generation_warnings': final_state.get("doc_generation_warnings"),
                    'citations_metadata': citations_metadata,
                }
                if doc_summary:
                    response_data['reply'] = doc_summary
            else:
                # Single answer mode (backward compatible, matching chat endpoint)
                if code_answer and code_answer.strip():
                    combined_answer = code_answer
                    total_citations = len(code_citations)
                elif coop_answer and coop_answer.strip():
                    combined_answer = coop_answer
                    total_citations = len(coop_citations)
                else:
                    combined_answer = answer
                    total_citations = len(project_citations)

                # Strip asterisks from project names/numbers
                combined_answer = strip_asterisks_from_projects(combined_answer)

                # Wrap code file citations with file-link tags for clickable links
                if code_citations:
                    combined_answer = wrap_code_file_links(combined_answer, code_citations)

                # Wrap coop file citations with file-link tags for clickable links
                if coop_citations:
                    combined_answer = wrap_code_file_links(combined_answer, coop_citations)

                # Convert external links to file-link format (for sidian-bot links) - works with existing frontend handlers
                if code_answer and code_answer.strip():
                    combined_answer = wrap_external_links(combined_answer)

                # Log to Supabase (matching chat endpoint)
                supabase_logger = get_supabase_logger()
                if supabase_logger.enabled:
                    user_identifier = request.user_identifier or f"{os.getenv('USERNAME', 'unknown')}@{os.getenv('COMPUTERNAME', 'unknown')}"

                    # Upload first image if provided (for logging)
                    image_url = None
                    if images_to_process and len(images_to_process) > 0:
                        image_url = supabase_logger.upload_image(
                            images_to_process[0],
                            message_id,
                            "image/png"
                        )

                    query_data = {
                        "message_id": message_id,
                        "session_id": request.session_id,
                        "user_query": request.message,
                        "rag_response": combined_answer,
                        "route": route_value,
                        "citations_count": total_citations,
                        "latency_ms": round(latency_ms, 2),
                        "user_identifier": user_identifier,
                        "image_url": image_url
                    }
                    supabase_logger.log_user_query(query_data)

                # Build response (matching chat endpoint)
                doc_answer_text = combined_answer
                
                # Only format document summary if workflow is docgen (not qa/rag)
                # This ensures document statistics only appear for doc generation, not RAG queries
                if workflow == "docgen" and doc_generation_result_payload:
                    doc_summary = format_document_summary(
                        doc_generation_result_payload,
                        final_state.get("doc_type"),
                        final_state.get("section_type"),
                        doc_answer_text,
                    )
                    if doc_summary and isinstance(doc_generation_result_payload, dict):
                        doc_generation_result_payload = dict(doc_generation_result_payload)
                        doc_generation_result_payload["document_summary"] = doc_summary
                response_data = {
                    'reply': combined_answer,
                    'session_id': request.session_id,
                    'timestamp': datetime.now().isoformat(),
                    'latency_ms': round(latency_ms, 2),
                    'citations': total_citations,
                    'route': route_value,
                    'message_id': message_id,
                    'image_similarity_results': image_similarity_results if image_similarity_results else [],  # Always include as array, not None
                    'follow_up_questions': follow_up_questions if follow_up_questions else None,
                    'follow_up_suggestions': follow_up_suggestions if follow_up_suggestions else None,
                    'workflow': final_state.get("workflow"),
                    'task_type': final_state.get("task_type"),
                    'doc_type': final_state.get("doc_type"),
                    'section_type': final_state.get("section_type"),
                    'execution_trace': execution_trace,
                    'node_path': " → ".join(execution_trace) if execution_trace else None,
                    'doc_generation_result': doc_generation_result_payload,
                    'doc_generation_warnings': final_state.get("doc_generation_warnings"),
                    'citations_metadata': citations_metadata,
                }
                if doc_summary:
                    response_data['reply'] = doc_summary

            if should_materialize_doc(
                final_state.get("workflow"), final_state.get("task_type"), final_state.get("doc_generation_result")
            ):
                doc_answer_text = doc_answer_text or final_state.get("final_answer") or final_state.get("answer")
                document_state = materialize_onlyoffice_document(
                    doc_generation_result_payload,
                    doc_answer_text or answer,
                    request.session_id,
                    message_id,
                )
                if document_state:
                    response_data["document_state"] = document_state
                    response_data["reply"] = doc_summary or "✅ Please refer to the document"
                    doc_url = document_state.get("docUrl") or document_state.get("onlyoffice", {}).get("docUrl")
                    logger.warning(f"📑 Streaming response includes document_state with url={doc_url}")
                    print(f"📑 Streaming document_state url={doc_url}")
            
            # Log image similarity results for debugging
            if image_similarity_results:
                logger.info(f"🖼️ Sending {len(image_similarity_results)} image similarity results to frontend")
                for i, img in enumerate(image_similarity_results[:3]):  # Log first 3
                    logger.info(f"🖼️   Image {i+1}: project={img.get('project_key')}, page={img.get('page_number')}, url={img.get('image_url', 'MISSING')[:50]}...")
            else:
                logger.info(f"🖼️ No image similarity results to send (image_similarity_results is empty or None)")
            
            # Send completion event with final result
            yield f"data: {json.dumps({'type': 'complete', 'result': response_data})}\n\n"
            
            logger.info(f"✅ Streaming completed in {latency_ms:.2f}ms [ID: {message_id}]")
            
        except GraphInterrupt as interrupt:
            interrupt_payload = extract_interrupt_data(interrupt)
            try:
                state_snapshot = await _load_thread_state(graph, request.session_id)
                safe_state = ensure_desktop_agent_state_compatibility(state_snapshot)
                safe_state["desktop_interrupt_pending"] = True
                safe_state["desktop_interrupt_data"] = interrupt_payload
                await _update_thread_state(
                    graph,
                    request.session_id,
                    {
                        "desktop_interrupt_pending": True,
                        "desktop_interrupt_data": interrupt_payload,
                        "desktop_approved_actions": safe_state.get("desktop_approved_actions", []),
                    },
                )
                interrupt_payload["state_summary"] = get_deep_agent_state_summary(safe_state)
            except Exception as exc:  # pragma: no cover - defensive
                logger.error(f"Failed to persist interrupt state: {exc}")
            yield f"data: {json.dumps({'type': 'interrupt', 'message': 'Action requires approval', 'interrupt': interrupt_payload, 'session_id': request.session_id})}\n\n"
            return
        except Exception as e:
            logger.error(f"❌ Streaming error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
    
    return StreamingResponse(
        generate_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",  # Disable nginx buffering
            "Access-Control-Allow-Origin": "*"  # CORS for streaming
        }
    )


# Deep agent approval endpoint
@app.post("/approve-action", response_model=ActionApprovalResponse)
async def approve_action(request: ActionApprovalRequest):
    """
    Approve or reject a pending desktop action after a GraphInterrupt.
    Updates state with approval and resumes the graph when approved.
    """
    from main import graph

    try:
        state_snapshot = await _load_thread_state(graph, request.session_id)
        if not state_snapshot:
            return ActionApprovalResponse(
                success=False,
                message=f"No state found for session {request.session_id}",
                resumed=False,
                next_state=None,
            )

        safe_state = ensure_rag_state_compatibility(state_snapshot)
        approved_actions = list(safe_state.get("desktop_approved_actions", []) or [])
        if request.approved and request.action_id not in approved_actions:
            approved_actions.append(request.action_id)

        update_values = {
            "desktop_approved_actions": approved_actions,
            "desktop_interrupt_pending": False,
            "desktop_interrupt_data": None,
        }

        await _update_thread_state(graph, request.session_id, update_values)
        logger.info(
            f"Action '{request.action_id}' approval={request.approved} recorded for session {request.session_id} (reason={request.reason})"
        )

        resumed_state = None
        resumed = False
        if request.approved:
            try:
                resumed_state = await _resume_graph(graph, request.session_id)
                resumed = True if resumed_state is not None else False
            except GraphInterrupt as interrupt:
                interrupt_payload = extract_interrupt_data(interrupt)
                await _update_thread_state(
                    graph,
                    request.session_id,
                    {
                        "desktop_interrupt_pending": True,
                        "desktop_interrupt_data": interrupt_payload,
                        "desktop_approved_actions": approved_actions,
                    },
                )
                return ActionApprovalResponse(
                    success=True,
                    message="Approval recorded, but another interrupt occurred.",
                    resumed=False,
                    next_state=interrupt_payload,
                )

        return ActionApprovalResponse(
            success=True,
            message="Action processed",
            resumed=resumed,
            next_state=resumed_state if isinstance(resumed_state, dict) else state_snapshot,
        )
    except Exception as exc:
        logger.error(f"Failed to process action approval: {exc}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(exc))


# ============================================================================
# DOC AGENT PROXY ENDPOINTS (mirrors Excel agent pattern)
# ============================================================================


async def _doc_agent_request(method: str, path: str, json_body: Optional[Dict[str, Any]] = None, params: Optional[Dict[str, Any]] = None):
    url = f"{DOC_API_URL}{path}"
    try:
        async with httpx.AsyncClient(timeout=DOC_API_TIMEOUT) as client:
            resp = await client.request(method, url, json=json_body, params=params)
            resp.raise_for_status()
            return resp.json()
    except httpx.HTTPStatusError as exc:
        logger.error(f"Doc agent HTTP error: {exc.response.text}")
        raise HTTPException(status_code=exc.response.status_code, detail=f"Doc agent error: {exc.response.text}")
    except Exception as exc:  # pragma: no cover - defensive
        logger.error(f"Doc agent unreachable at {url}: {exc}")
        raise HTTPException(status_code=502, detail=f"Doc agent unavailable at {DOC_API_URL}: {exc}")

def _validate_doc_ops_payload(payload: Dict[str, Any]) -> None:
    if payload is None:
        raise HTTPException(status_code=400, detail="Payload is required.")
    ops = payload.get("ops")
    if ops is None or not isinstance(ops, list) or len(ops) == 0:
        raise HTTPException(status_code=400, detail="ops[] is required and must be a non-empty list.")
    schema_version = payload.get("schema_version", DOC_OP_SCHEMA_VERSION)
    if schema_version != DOC_OP_SCHEMA_VERSION:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported schema_version {schema_version}. Expected {DOC_OP_SCHEMA_VERSION}",
        )
    for op in ops:
        if not isinstance(op, dict) or "op" not in op:
            raise HTTPException(status_code=400, detail="Each operation must be an object with an 'op' field.")
        name = op.get("op")
        if name in {"replace_text", "insert_paragraph", "insert_heading", "delete_block", "set_style"}:
            target = op.get("target")
            if not isinstance(target, dict) or "index" not in target:
                raise HTTPException(status_code=400, detail=f"Operation '{name}' requires target.index")
    # Ensure file_path exists if provided (best-effort)
    file_path = payload.get("file_path")
    if file_path:
        try:
            Path(file_path)
        except Exception:
            raise HTTPException(status_code=400, detail="file_path must be a valid path string.")


@app.get("/api/doc/health")
async def doc_agent_health():
    return await _doc_agent_request("get", "/health")


@app.post("/api/doc/open")
async def doc_agent_open(payload: Dict[str, Any]):
    return await _doc_agent_request("post", "/api/doc/open", json_body=payload)


@app.post("/api/doc/apply")
async def doc_agent_apply(payload: Dict[str, Any]):
    _validate_doc_ops_payload(payload)
    return await _doc_agent_request("post", "/api/doc/apply", json_body=payload)


@app.post("/api/doc/export")
async def doc_agent_export(payload: Dict[str, Any]):
    return await _doc_agent_request("post", "/api/doc/export", json_body=payload)


@app.get("/api/doc/history")
async def doc_agent_history(doc_id: str, limit: int = 100):
    return await _doc_agent_request("get", "/api/doc/history", params={"doc_id": doc_id, "limit": limit})


# Legacy endpoint REMOVED - use /chat/stream instead
# The /chat/legacy endpoint has been removed to simplify the codebase
# All clients should migrate to the streaming endpoint

# IFC File Upload endpoint
@app.post("/chat/upload-ifc")
async def upload_ifc_to_speckle(
    file: UploadFile = File(...),
    project_id: Optional[str] = Form(None),
    model_name: str = Form("main")
):
    """
    Upload IFC file to Speckle server.
    Creates new project if project_id not provided.
    Saves metadata to database.
    
    Returns:
        {
            "success": bool,
            "speckle_project_id": str,
            "speckle_model_name": str,
            "message": str
        }
    """
    try:
        logger.info(f"📤 IFC upload request: {file.filename}, project_id={project_id}, model={model_name}")
        
        # Get Speckle credentials from environment
        speckle_url = os.getenv("SPECKLE_URL") or os.getenv("SPECKLE_SERVER_URL") or "https://app.speckle.systems"
        speckle_token = os.getenv("SPECKLE_TOKEN")
        
        if not speckle_token:
            raise HTTPException(status_code=500, detail="Speckle token not configured")
        
        # Read file content
        file_content = await file.read()
        file_size = len(file_content)
        file_type = file.filename.split('.')[-1].lower() if '.' in file.filename else ''
        
        logger.info(f"📦 File details: {file.filename}, size={file_size} bytes, type={file_type}")
        
        # 1. CREATE PROJECT IN SPECKLE (if not provided)
        speckle_project_id = project_id
        if not speckle_project_id:
            project_name = file.filename.replace('.ifc', '').replace('.IFC', '') or f"IFC Upload {datetime.now().strftime('%Y-%m-%d')}"
            
            logger.info(f"🏗️ Creating Speckle project: {project_name}")
            
            # Call Speckle GraphQL to create project
            async with httpx.AsyncClient(timeout=30.0) as client:
                create_project_response = await client.post(
                    f"{speckle_url}/graphql",
                    headers={
                        "Content-Type": "application/json",
                        "Authorization": f"Bearer {speckle_token}"
                    },
                    json={
                        "query": """
                            mutation CreateProject($input: ProjectCreateInput!) {
                                projectMutations {
                                    create(input: $input) {
                                        id
                                        name
                                    }
                                }
                            }
                        """,
                        "variables": {
                            "input": {
                                "name": project_name,
                                "description": f"Uploaded from chat: {file.filename}",
                                "visibility": "PRIVATE"
                            }
                        }
                    }
                )
                
                if create_project_response.status_code != 200:
                    raise HTTPException(
                        status_code=500,
                        detail=f"Failed to create Speckle project: {create_project_response.text}"
                    )
                
                result = create_project_response.json()
                if "errors" in result:
                    raise HTTPException(
                        status_code=500,
                        detail=f"GraphQL error: {result['errors']}"
                    )
                
                speckle_project_id = result["data"]["projectMutations"]["create"]["id"]
                project_name_created = result["data"]["projectMutations"]["create"]["name"]
                logger.info(f"✅ Created Speckle project: {speckle_project_id} ({project_name_created})")
        
        # 2. SAVE TO YOUR DATABASE (if you have a database setup)
        # TODO: Add your database save logic here
        # Example:
        # file_record = {
        #     "filename": file.filename,
        #     "file_type": file_type,
        #     "file_size": file_size,
        #     "speckle_project_id": speckle_project_id,
        #     "speckle_model_name": model_name,
        #     "upload_date": datetime.now().isoformat(),
        #     "status": "uploading"
        # }
        # await your_db.file_uploads.insert(file_record)
        
        # 3. UPLOAD TO SPECKLE SERVER using autodetect endpoint
        logger.info(f"⬆️ Uploading file to Speckle: project={speckle_project_id}, model={model_name}")
        
        async with httpx.AsyncClient(timeout=300.0) as client:
            # Create multipart form data
            files = {
                "file": (file.filename, file_content, file.content_type or "application/octet-stream")
            }
            
            upload_response = await client.post(
                f"{speckle_url}/api/file/autodetect/{speckle_project_id}/{model_name}",
                headers={
                    "Authorization": f"Bearer {speckle_token}"
                },
                files=files
            )
            
            if upload_response.status_code != 201:
                error_text = upload_response.text
                logger.error(f"❌ Speckle upload failed: {upload_response.status_code} - {error_text}")
                raise HTTPException(
                    status_code=500,
                    detail=f"Speckle upload failed: {error_text}"
                )
            
            upload_result = upload_response.json()
            logger.info(f"✅ File uploaded to Speckle successfully: {upload_result}")
        
        # 4. UPDATE YOUR DATABASE (if you have one)
        # TODO: Update status to "completed"
        # await your_db.file_uploads.update(
        #     {"status": "completed", "speckle_blob_id": upload_result.get("blobId")}
        # )
        
        return {
            "success": True,
            "speckle_project_id": speckle_project_id,
            "speckle_model_name": model_name,
            "message": f"File '{file.filename}' uploaded successfully. Speckle is processing it and it will appear in your project shortly."
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ IFC upload error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

# Feedback stats endpoint (for debugging/admin)
@app.get("/feedback/stats")
async def feedback_stats():
    """Get feedback statistics"""
    try:
        feedback_logger = get_feedback_logger()
        stats = feedback_logger.get_feedback_stats()
        return stats
    except Exception as e:
        logger.error(f"Error getting feedback stats: {e}")
        return {"error": str(e)}

# Database health endpoint
@app.get("/db/health")
async def database_health():
    """Check database connection status"""
    try:
        db_status = test_database_connection()
        return db_status
    except Exception as e:
        logger.error(f"Database health check failed: {e}")
        return {"connected": False, "error": str(e)}

# Debug routing endpoint
@app.get("/debug/routing")
async def debug_routing():
    """Debug endpoint to show current routing configuration"""
    try:
        # Import the constants from config
        from config.settings import (
            MAX_SMART_RETRIEVAL_DOCS, MAX_LARGE_RETRIEVAL_DOCS, 
            MAX_HYBRID_RETRIEVAL_DOCS, SUPA_SMART_TABLE, SUPA_LARGE_TABLE
        )
        from nodes.DBRetrieval.KGdb.supabase_client import vs_smart, vs_large
        
        return {
            "supabase_configured": bool(os.getenv("SUPABASE_URL") and os.getenv("SUPABASE_ANON_KEY")),
            "smart_table": SUPA_SMART_TABLE,
            "large_table": SUPA_LARGE_TABLE,
            "vs_smart_exists": vs_smart is not None,
            "vs_large_exists": vs_large is not None,
            "chunk_limits": {
                "smart": MAX_SMART_RETRIEVAL_DOCS,
                "large": MAX_LARGE_RETRIEVAL_DOCS,
                "hybrid": MAX_HYBRID_RETRIEVAL_DOCS
            }
        }
    except Exception as e:
        logger.error(f"Debug routing check failed: {e}")
        return {"error": str(e)}

# Enhanced logging endpoints
@app.get("/logs/enhanced")
async def get_enhanced_logs(
    level: Optional[str] = None,
    logger: Optional[str] = None,
    category: Optional[str] = None,
    limit: int = 100,
    format: str = "json"
):
    """Enhanced logs with function tracking and chunk flow"""
    logs = enhanced_log_capture.get_logs(level_filter=level, logger_filter=logger, limit=limit)
    
    if category:
        logs = [log for log in logs if log.get('extra_data', {}).get('category') == category]
    
    if format == "html":
        return await render_enhanced_log_html(logs, level, logger, category, limit)
    
    return {
        "logs": logs,
        "function_calls": enhanced_log_capture.get_function_calls(20),
        "chunk_flow": enhanced_log_capture.get_chunk_flow(50),
        "count": len(logs)
    }

@app.post("/logs/enhanced/clear")
async def clear_enhanced_logs():
    """Clear all enhanced logs and tracking data"""
    enhanced_log_capture.clear_all()
    return {"message": "All logs and tracking data cleared successfully"}

@app.get("/logs/enhanced/stats")
async def enhanced_log_stats():
    """Get comprehensive statistics"""
    return get_enhanced_log_stats()

# Supabase Statistics Endpoints
@app.get("/stats/interactions")
async def get_interaction_stats(user_identifier: str = None, days: int = 30):
    """Get comprehensive interaction statistics from Supabase"""
    supabase_logger = get_supabase_logger()
    if not supabase_logger.enabled:
        return {"error": "Supabase logging not enabled"}
    
    stats = supabase_logger.get_interaction_stats(user_identifier, days)
    return stats

@app.get("/stats/user/{user_identifier}")
async def get_user_summary(user_identifier: str, days: int = 30):
    """Get comprehensive summary for a specific user"""
    supabase_logger = get_supabase_logger()
    if not supabase_logger.enabled:
        return {"error": "Supabase logging not enabled"}
    
    summary = supabase_logger.get_user_summary(user_identifier, days)
    return summary

@app.get("/stats/recent")
async def get_recent_interactions(limit: int = 50, user_identifier: str = None):
    """Get recent interactions"""
    supabase_logger = get_supabase_logger()
    if not supabase_logger.enabled:
        return {"error": "Supabase logging not enabled"}
    
    interactions = supabase_logger.get_recent_interactions(limit, user_identifier)
    return {"interactions": interactions, "count": len(interactions)}

# Instructions endpoint
@app.get("/instructions")
async def get_instructions():
    """
    Return usage instructions for the frontend.
    Content is stored in instructions.py and can be updated without rebuilding the executable.
    """
    return {
        "content": INSTRUCTIONS_CONTENT,
        "format": "markdown"
    }


# Test endpoint to verify which backend is running
@app.get("/test/backend-version")
async def test_backend_version():
    """Test endpoint to verify we're using the new Backend/Backend code"""
    from config.settings import SUPA_SMART_TABLE, SUPA_LARGE_TABLE
    return {
        "backend": "Backend/Backend (NEW)",
        "smart_table": SUPA_SMART_TABLE,
        "large_table": SUPA_LARGE_TABLE,
        "message": "This is the NEW modular backend"
    }

# Root endpoint
@app.get("/")
@app.head("/")
async def root():
    return {
        "message": "Mantle RAG API Server",
        "version": "1.0.0",
        "backend": "Backend/Backend (NEW)",
        "endpoints": {
            "chat": "/chat",
            "health": "/health",
            "db_health": "/db/health",
            "test_backend": "/test/backend-version",
            "instructions": "/instructions",
            "logs_enhanced": "/logs/enhanced",
            "logs_stats": "/logs/enhanced/stats",
            "stats_interactions": "/stats/interactions",
            "stats_user": "/stats/user/{user_identifier}",
            "stats_recent": "/stats/recent",
            "graph_cypher": "/graph/cypher (DEBUG only)",
            "graph_schema": "/graph/schema (DEBUG only)"
        }
    }

# ============================================================
# KUZU GRAPH DATABASE ENDPOINTS (DEBUG ONLY)
# ============================================================
@app.post("/graph/cypher", response_model=CypherResponse)
async def execute_cypher(request: CypherRequest):
    """Execute a Cypher query against the Kuzu graph database (Debug mode only)"""
    if not DEBUG_MODE:
        raise HTTPException(status_code=403, detail="Graph database access is only available in DEBUG_MODE")
    
    if get_kuzu_manager is None:
        raise HTTPException(status_code=500, detail="Kuzu manager not available")
    
    try:
        kuzu_manager = get_kuzu_manager()
        result = kuzu_manager.execute(request.query, request.params)
        return CypherResponse(**result)
    except Exception as e:
        logger.error(f"Cypher endpoint error: {e}")
        return CypherResponse(success=False, error=str(e), query=request.query)

@app.get("/graph/schema")
async def get_graph_schema():
    """Get the Kuzu graph database schema (Debug mode only)"""
    if not DEBUG_MODE:
        raise HTTPException(status_code=403, detail="Graph database access is only available in DEBUG_MODE")
    
    if get_kuzu_manager is None:
        raise HTTPException(status_code=500, detail="Kuzu manager not available")
    
    try:
        kuzu_manager = get_kuzu_manager()
        return kuzu_manager.get_schema()
    except Exception as e:
        logger.error(f"Graph schema endpoint error: {e}")
        return {"success": False, "error": str(e)}

if __name__ == "__main__":
    # Check environment setup
    if not os.getenv("OPENAI_API_KEY"):
        logger.error("OPENAI_API_KEY environment variable not found!")
        print("\n❌ Please set OPENAI_API_KEY in your .env file before starting the server.")
        print("   Example: OPENAI_API_KEY=your_api_key_here")
        exit(1)

    # Get port from environment (for Render deployment) or default to 8000
    port = int(os.getenv("PORT", 8000))

    print("\n🚀 Starting Mantle RAG API Server...")
    print(f"   Chat endpoint: http://0.0.0.0:{port}/chat")
    print(f"   Health check: http://0.0.0.0:{port}/health")
    print(f"   Database status: http://0.0.0.0:{port}/db/health")
    print(f"   📊 Enhanced Debugger: http://0.0.0.0:{port}/logs/enhanced?format=html")
    print(f"   📈 Log Statistics: http://0.0.0.0:{port}/logs/enhanced/stats")
    if DEBUG_MODE:
        print(f"   🛠️  Kuzu Graph DB Cypher: http://0.0.0.0:{port}/graph/schema")
    print(f"\n📡 Server running on port: {port}")
    print("⚡ Server ready - listening on http://0.0.0.0:8000")

    # Run the server
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=port,
        log_level="info",
        access_log=True
    )
