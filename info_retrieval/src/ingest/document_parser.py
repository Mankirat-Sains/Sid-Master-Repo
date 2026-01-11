from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

from utils.logger import get_logger

logger = get_logger(__name__)

try:
    import docx  # type: ignore
    from docx.opc.constants import RELATIONSHIP_TYPE  # type: ignore
except ImportError:  # pragma: no cover - exercised when dependency is missing
    docx = None
    RELATIONSHIP_TYPE = None

try:
    import fitz  # PyMuPDF # type: ignore
except ImportError:  # pragma: no cover - exercised when dependency is missing
    fitz = None


@dataclass
class Section:
    title: str
    content: str
    level: int = 1
    children: List["Section"] = field(default_factory=list)


@dataclass
class TableData:
    rows: List[List[str]] = field(default_factory=list)


@dataclass
class ParsedDocument:
    text: str
    sections: List[Section] = field(default_factory=list)
    tables: List[TableData] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    pages: List[Dict[str, Any]] = field(default_factory=list)
    artifact_id: Optional[str] = None
    version_id: Optional[str] = None


def parse_docx(file_path: str | Path, company_id: Optional[str] = None, source: str = "upload") -> ParsedDocument:
    """
    Parse a .docx file into a structured document representation.
    """
    if docx is None:
        raise ImportError("python-docx is required to parse .docx files.")

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    logger.info("Parsing DOCX file %s", path)
    try:
        document = docx.Document(path)  # type: ignore
    except Exception as exc:
        logger.error("Failed to open DOCX %s: %s", path, exc)
        raise

    text_parts: List[str] = []
    sections: List[Section] = []
    tables: List[TableData] = []

    stack: List[Section] = []

    for para in document.paragraphs:
        paragraph_text = para.text.strip()
        if not paragraph_text:
            continue

        text_parts.append(paragraph_text)
        style_name = para.style.name if para.style else ""
        heading_level = _extract_heading_level(style_name)
        if heading_level is not None:
            new_section = Section(title=paragraph_text, content="", level=heading_level)
            while stack and stack[-1].level >= heading_level:
                stack.pop()
            if stack:
                stack[-1].children.append(new_section)
            else:
                sections.append(new_section)
            stack.append(new_section)
        elif stack:
            stack[-1].content = _append_paragraph(stack[-1].content, paragraph_text)
        else:
            # No heading encountered yet; start a default section
            default_section = Section(title="Introduction", content=paragraph_text, level=1)
            sections.append(default_section)
            stack.append(default_section)

    for table in document.tables:
        rows: List[List[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(TableData(rows=rows))

    headers = _extract_header_footer(document, header=True)
    footers = _extract_header_footer(document, header=False)
    core_props = _extract_core_properties(document)
    images = _extract_image_rels(document)
    hyperlinks = _extract_hyperlinks(document)

    full_text = "\n".join(text_parts)
    artifact_id = generate_artifact_id(path, company_id) if company_id else None
    version_id = generate_version_id(full_text)
    metadata = {
        "file_path": str(path),
        "company_id": company_id,
        "source": source,
        "page_count": None,
        "headers": headers,
        "footers": footers,
        "core_properties": core_props,
        "image_count": images,
        "hyperlinks": hyperlinks,
    }
    return ParsedDocument(
        text=full_text,
        sections=sections,
        tables=tables,
        metadata=metadata,
        artifact_id=artifact_id,
        version_id=version_id,
    )


def parse_pdf(file_path: str | Path, company_id: Optional[str] = None, source: str = "upload") -> ParsedDocument:
    """
    Parse a PDF into text and inferred sections using PyMuPDF.
    Gracefully skips unreadable pages and logs errors.
    """
    if fitz is None:
        raise ImportError("PyMuPDF (fitz) is required to parse PDF files.")

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    logger.info("Parsing PDF with PyMuPDF: %s", path)
    doc = fitz.open(str(path))  # type: ignore
    pages: List[Dict[str, Any]] = []
    text_parts: List[str] = []

    for page_index, page in enumerate(doc):
        try:
            text = page.get_text("text")
            text_parts.append(text or "")
            blocks = page.get_text("blocks")
            pages.append(
                {
                    "page_number": page_index + 1,
                    "text": text,
                    "blocks": blocks,
                }
            )
        except Exception as exc:  # pragma: no cover - defensive
            logger.error("Failed to parse page %s in %s: %s", page_index + 1, path, exc)
            continue

    full_text = "\n".join(text_parts)
    sections = infer_sections_from_text(full_text)
    artifact_id = generate_artifact_id(path, company_id) if company_id else None
    version_id = generate_version_id(full_text)
    metadata: Dict[str, Any] = {
        "file_path": str(path),
        "company_id": company_id,
        "source": source,
        "page_count": doc.page_count,
        "author": doc.metadata.get("author", ""),
        "title": doc.metadata.get("title", ""),
        "creation_date": doc.metadata.get("creationDate", ""),
    }
    return ParsedDocument(
        text=full_text,
        sections=sections,
        tables=[],
        metadata=metadata,
        pages=pages,
        artifact_id=artifact_id,
        version_id=version_id,
    )


def extract_sections(document: ParsedDocument) -> List[Section]:
    """
    Return existing sections or infer them from document text.
    """
    if document.sections:
        return document.sections
    return infer_sections_from_text(document.text)


def extract_tables(document: ParsedDocument) -> List[TableData]:
    return document.tables


def infer_sections_from_text(text: str) -> List[Section]:
    """
    Heuristic section detection from raw text.
    Treat numbered headings or fully capitalized lines as section titles.
    """
    # Accept headings like "1. Introduction", "2.1 Scope", or fully capitalized lines.
    heading_pattern = re.compile(r"^(\d+(?:\.\d+)*\.?\s+.+|[A-Z][A-Z\s]{3,})$")
    sections: List[Section] = []
    current_section: Section | None = None

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        if heading_pattern.match(stripped):
            if current_section:
                sections.append(current_section)
            current_section = Section(title=stripped, content="", level=_infer_level_from_heading(stripped))
        elif current_section:
            current_section.content = _append_paragraph(current_section.content, stripped)
        else:
            current_section = Section(title="Introduction", content=stripped, level=1)

    if current_section:
        sections.append(current_section)

    return sections


def generate_artifact_id(file_path: str | Path, company_id: Optional[str] = None) -> str:
    """
    Generate a stable artifact identifier using company + filename + file timestamp.
    """
    path = Path(file_path)
    stem = path.stem
    prefix = f"{company_id}_" if company_id else ""
    try:
        timestamp = path.stat().st_mtime
    except FileNotFoundError:
        timestamp = 0
    base = f"{company_id or 'org'}:{stem}:{timestamp}"
    digest = hashlib.sha256(base.encode("utf-8")).hexdigest()[:10]
    return f"{prefix}{stem}_{digest}"


def generate_version_id(document_content: str) -> str:
    """
    Version hash based on document content to detect changes.
    """
    return hashlib.sha256(document_content.encode("utf-8")).hexdigest()


def _append_paragraph(existing: str, new_text: str) -> str:
    return f"{existing}\n{new_text}" if existing else new_text


def _extract_heading_level(style_name: str) -> int | None:
    match = re.match(r"Heading\s+(\d+)", style_name)
    if match:
        return int(match.group(1))
    return None


def _extract_header_footer(document: Any, header: bool = True) -> List[str]:
    texts: List[str] = []
    try:
        for sect in document.sections:  # type: ignore
            hf = sect.header if header else sect.footer
            for para in hf.paragraphs:
                if para.text.strip():
                    texts.append(para.text.strip())
    except Exception:
        return []
    return texts


def _extract_core_properties(document: Any) -> Dict[str, Any]:
    props: Dict[str, Any] = {}
    try:
        cp = document.core_properties  # type: ignore
        props = {
            "author": cp.author,
            "created": getattr(cp, "created", None),
            "modified": getattr(cp, "modified", None),
            "title": cp.title,
            "subject": cp.subject,
            "comments": cp.comments,
        }
    except Exception:
        return {}
    return props


def _extract_image_rels(document: Any) -> int:
    try:
        rels = document.part.related_parts.values()  # type: ignore
        return sum(1 for rel in rels if getattr(rel, "content_type", "").startswith("image/"))
    except Exception:
        return 0


def _extract_hyperlinks(document: Any) -> List[str]:
    links: List[str] = []
    if RELATIONSHIP_TYPE is None:
        return links
    try:
        rels = document.part.rels  # type: ignore
        for rel in rels.values():
            if rel.reltype == RELATIONSHIP_TYPE.HYPERLINK:
                links.append(rel.target_ref)
    except Exception:
        return []
    return links


def _infer_level_from_heading(heading: str) -> int:
    numeric = re.match(r"(\d+)", heading)
    if numeric:
        return len(numeric.group(1).split("."))
    return 1
