from pathlib import Path

from fastapi.testclient import TestClient
from docx import Document

import doc_agent_api


def test_apply_ops_end_to_end(tmp_path: Path):
    # Reset globals for isolated test
    doc_agent_api._agent_config = None
    doc_agent_api._doc_agent = None
    doc_agent_api._history = {}

    workspace = tmp_path / "workspace"
    workspace.mkdir(parents=True, exist_ok=True)

    doc_path = workspace / "test.docx"
    doc = Document()
    doc.add_paragraph("Intro text")
    doc.add_paragraph("Body text")
    doc.save(doc_path)

    client = TestClient(doc_agent_api.app)

    config_payload = {
        "workspace_dir": str(workspace),
        "docs": [
            {"doc_id": "test-doc", "file_path": str(doc_path)}
        ],
    }

    res_conf = client.post("/api/doc/configure", json=config_payload)
    assert res_conf.status_code == 200

    apply_payload = {
        "schema_version": 1,
        "file_path": str(doc_path),
        "ops": [
            {"op": "replace_text", "target": {"index": 0}, "text": "Updated intro"},
            {"op": "insert_heading", "target": {"index": 1}, "level": 2, "text": "Schedule"},
        ],
    }

    res_apply = client.post("/api/doc/apply", json=apply_payload)
    assert res_apply.status_code == 200
    data = res_apply.json()
    assert data["doc_id"] == "test-doc"
    assert "change_summary" in data
    assert any("Updated intro" in blk["text"] for blk in data["structure"]["blocks"])
