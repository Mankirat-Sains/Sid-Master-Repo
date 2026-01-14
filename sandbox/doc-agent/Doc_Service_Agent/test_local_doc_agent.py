import tempfile
from pathlib import Path

from docx import Document

from local_doc_agent import DocAgent


def test_apply_ops_replace_and_insert(tmp_path: Path):
    doc_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Intro")
    doc.add_paragraph("Body")
    doc.save(doc_path)

    agent = DocAgent({"workspace_dir": str(tmp_path), "docs": [{"doc_id": "sample", "file_path": str(doc_path)}]})
    result = agent.apply_ops(
        doc_id="sample",
        file_path=None,
        ops=[
          {"op": "replace_text", "target": {"index": 0}, "text": "Updated intro"},
          {"op": "insert_heading", "target": {"index": 1}, "level": 2, "text": "Schedule"},
        ],
    )

    assert result["doc_id"] == "sample"
    assert any("Updated intro" in blk["text"] for blk in result["structure"]["blocks"])
    assert any(blk["type"] == "heading" for blk in result["structure"]["blocks"])
