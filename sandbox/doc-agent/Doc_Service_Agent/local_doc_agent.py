"""
Local DOCX editing agent.
Provides deterministic operations (insert/replace/delete/style/reorder) over DOCX files.
"""

from __future__ import annotations

import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


class DocOperationError(Exception):
    """Raised when an operation cannot be applied."""


def load_config(path: str) -> Dict[str, Any]:
    cfg_path = Path(path)
    if not cfg_path.exists():
        raise FileNotFoundError(f"Config not found: {cfg_path}")
    data = json.loads(cfg_path.read_text())
    return data


class DocAgent:
    """Manages DOCX files and applies deterministic operations."""

    def __init__(self, config: Optional[Dict[str, Any]] = None) -> None:
        self.config = config or {}
        self.workspace = Path(self.config.get("workspace_dir", "./workspace/doc-agent")).resolve()
        self.workspace.mkdir(parents=True, exist_ok=True)
        self.docs: Dict[str, Path] = {}
        self.history: Dict[str, List[Dict[str, Any]]] = {}

        for doc_cfg in self.config.get("docs", []):
            doc_id = doc_cfg.get("doc_id") or Path(doc_cfg.get("file_path", "")).stem
            doc_path = self._resolve_path(doc_cfg.get("file_path", ""))
            if doc_id and doc_path:
                self.docs[doc_id] = doc_path
                logger.info(f"Registered doc '{doc_id}' -> {doc_path}")

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------
    def open_doc(self, file_path: str, doc_id: Optional[str] = None) -> Dict[str, Any]:
        doc_path = self._resolve_path(file_path)
        if not doc_path.exists():
            raise DocOperationError(f"File does not exist: {doc_path}")

        doc_id = doc_id or doc_path.stem
        self.docs[doc_id] = doc_path
        structure = self._serialize_document(Document(doc_path))
        return {"doc_id": doc_id, "file_path": str(doc_path), "structure": structure}

    def apply_ops(
        self,
        doc_id: Optional[str],
        file_path: Optional[str],
        ops: List[Dict[str, Any]],
        save_as: Optional[str] = None,
    ) -> Dict[str, Any]:
        target_id, doc_path = self._resolve_doc(doc_id, file_path)
        doc = self._load_or_create(doc_path)

        change_summary: List[str] = []
        for idx, op in enumerate(ops):
            success, msg = self._apply_single_op(doc, op)
            change_summary.append(f"[{idx}] {msg}")
            if not success:
                raise DocOperationError(f"Operation failed: {msg}")

        out_path = self._persist_doc(doc, doc_path, save_as)
        structure = self._serialize_document(doc)

        entry = {
            "doc_id": target_id,
            "file_path": str(out_path),
            "change_summary": change_summary,
            "timestamp": datetime.utcnow().isoformat(),
        }
        self.history.setdefault(target_id, []).append(entry)
        self.history[target_id] = self.history[target_id][-200:]

        return {
            "doc_id": target_id,
            "file_path": str(out_path),
            "change_summary": change_summary,
            "structure": structure,
        }

    def export_doc(self, doc_id: str, target_path: Optional[str] = None) -> Dict[str, Any]:
        if doc_id not in self.docs:
            raise DocOperationError(f"Doc '{doc_id}' not registered.")

        src = self.docs[doc_id]
        if not src.exists():
            raise DocOperationError(f"Doc file not found: {src}")

        target = Path(target_path) if target_path else src
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_bytes(src.read_bytes())

        return {
            "doc_id": doc_id,
            "source_path": str(src),
            "target_path": str(target),
            "size_bytes": target.stat().st_size,
        }

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------
    def _resolve_doc(self, doc_id: Optional[str], file_path: Optional[str]) -> Tuple[str, Path]:
        if doc_id and doc_id in self.docs:
            return doc_id, self.docs[doc_id]

        if file_path:
            path = self._resolve_path(file_path)
            target_id = doc_id or path.stem
            self.docs[target_id] = path
            return target_id, path

        raise DocOperationError("doc_id or file_path must be provided.")

    def _resolve_path(self, file_path: str) -> Path:
        """Resolve a path to an absolute location anchored in the workspace when relative."""
        p = Path(file_path)
        if not p.is_absolute():
            p = (self.workspace / p).resolve()
        else:
            p = p.resolve()
        p.parent.mkdir(parents=True, exist_ok=True)
        return p

    def _load_or_create(self, doc_path: Path) -> Document:
        if doc_path.exists():
            return Document(doc_path)
        doc_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_paragraph("")  # seed doc
        return doc

    def _persist_doc(self, doc: Document, doc_path: Path, save_as: Optional[str]) -> Path:
        target = Path(save_as) if save_as else doc_path
        if not target.is_absolute():
            target = (self.workspace / target).resolve()
        target.parent.mkdir(parents=True, exist_ok=True)
        doc.save(target)
        return target

    def _serialize_document(self, doc: Document) -> Dict[str, Any]:
        blocks = []
        for idx, para in enumerate(doc.paragraphs):
            blocks.append(
                {
                    "index": idx,
                    "type": "heading" if para.style.name.startswith("Heading") else "paragraph",
                    "text": para.text,
                    "style": para.style.name if para.style else None,
                    "level": self._heading_level(para),
                }
            )
        return {"blocks": blocks, "paragraph_count": len(doc.paragraphs)}

    def _heading_level(self, para: Paragraph) -> Optional[int]:
        if para.style and para.style.name.startswith("Heading"):
            try:
                return int(para.style.name.replace("Heading", "").strip())
            except Exception:
                return None
        return None

    # ------------------------------------------------------------------
    # Operation execution
    # ------------------------------------------------------------------
    def _apply_single_op(self, doc: Document, op: Dict[str, Any]) -> Tuple[bool, str]:
        op_name = (op.get("op") or "").lower()
        target = op.get("target", {}) or {}

        if op_name == "replace_text":
            return self._op_replace_text(doc, target, op.get("text", ""))
        if op_name == "insert_paragraph":
            return self._op_insert_paragraph(doc, target, op.get("text", ""), op.get("style"))
        if op_name == "insert_heading":
            return self._op_insert_heading(doc, target, op.get("text", ""), op.get("level", 2))
        if op_name == "delete_block":
            return self._op_delete_block(doc, target)
        if op_name == "set_style":
            return self._op_set_style(doc, target, op.get("style"))
        if op_name == "reorder_blocks":
            return self._op_reorder_blocks(doc, op.get("from_index"), op.get("to_index"))

        return False, f"Unknown op '{op_name}'"

    def _get_para(self, doc: Document, target: Dict[str, Any]) -> Paragraph:
        idx = target.get("index")
        if idx is None or idx < 0 or idx >= len(doc.paragraphs):
            raise DocOperationError(f"Paragraph index out of range: {idx}")
        return doc.paragraphs[idx]

    def _op_replace_text(self, doc: Document, target: Dict[str, Any], text: str) -> Tuple[bool, str]:
        para = self._get_para(doc, target)
        para.text = text
        return True, f"replace_text paragraph {target.get('index')} -> {text[:60]}"

    def _op_insert_paragraph(
        self, doc: Document, target: Dict[str, Any], text: str, style: Optional[str]
    ) -> Tuple[bool, str]:
        insert_after = target.get("index")
        new_para = self._insert_after(doc, insert_after, text)
        if style:
            new_para.style = style
        return True, f"insert_paragraph after {insert_after}"

    def _op_insert_heading(
        self, doc: Document, target: Dict[str, Any], text: str, level: int
    ) -> Tuple[bool, str]:
        insert_after = target.get("index")
        new_para = self._insert_after(doc, insert_after, text)
        new_para.style = f"Heading {level}"
        return True, f"insert_heading after {insert_after} level {level}"

    def _insert_after(self, doc: Document, after_idx: Optional[int], text: str) -> Paragraph:
        if after_idx is None or after_idx < 0 or after_idx >= len(doc.paragraphs) - 1:
            return doc.add_paragraph(text)

        # Insert before the next paragraph to effectively place after current index
        next_para = doc.paragraphs[after_idx + 1]
        new_para = next_para.insert_paragraph_before(text)
        return new_para

    def _op_delete_block(self, doc: Document, target: Dict[str, Any]) -> Tuple[bool, str]:
        para = self._get_para(doc, target)
        p_element = para._element
        p_element.getparent().remove(p_element)
        para._p = None  # type: ignore[attr-defined]
        return True, f"delete_block index {target.get('index')}"

    def _op_set_style(self, doc: Document, target: Dict[str, Any], style: Optional[str]) -> Tuple[bool, str]:
        if not style:
            raise DocOperationError("set_style requires a style name.")
        para = self._get_para(doc, target)
        para.style = style
        return True, f"set_style index {target.get('index')} -> {style}"

    def _op_reorder_blocks(self, doc: Document, from_idx: Any, to_idx: Any) -> Tuple[bool, str]:
        if from_idx is None or to_idx is None:
            raise DocOperationError("reorder_blocks requires from_index and to_index.")
        if from_idx < 0 or from_idx >= len(doc.paragraphs):
            raise DocOperationError(f"from_index out of range: {from_idx}")
        if to_idx < 0 or to_idx > len(doc.paragraphs):
            raise DocOperationError(f"to_index out of range: {to_idx}")

        body = doc._body._element  # noqa: SLF001
        para_elem = doc.paragraphs[from_idx]._element
        body.remove(para_elem)
        body.insert(to_idx, para_elem)
        return True, f"reorder_blocks {from_idx} -> {to_idx}"
