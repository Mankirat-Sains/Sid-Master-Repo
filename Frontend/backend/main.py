"""
Simple FastAPI backend for SidOS Word Export and Agents Integration
"""
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional, List, Dict
import os
import sys
import tempfile
from pathlib import Path
from datetime import datetime
import time

# Load .env file BEFORE importing agents (agents need env vars at import time)
# This MUST happen before any agent imports because agents initialize Supabase at module level
try:
    from dotenv import load_dotenv
    # .env file is in Frontend folder (parent of backend)
    env_path = Path(__file__).parent.parent / ".env"
    if env_path.exists():
        # Load with override=True to ensure our values take precedence
        load_dotenv(env_path, override=True)
        print(f"✓ Loaded .env from {env_path}")
        # Verify critical env vars are loaded
        if os.getenv("OPENAI_API_KEY"):
            print("✓ OPENAI_API_KEY found")
        if os.getenv("SUPABASE_URL"):
            print("✓ SUPABASE_URL found")
        if os.getenv("SUPABASE_KEY") or os.getenv("SUPABASE_ANON_KEY"):
            print("✓ SUPABASE_KEY found")
    else:
        # Try loading from current directory as fallback
        load_dotenv(override=True)
        print(f"⚠ .env file not found at {env_path}, trying current directory")
        if not os.getenv("OPENAI_API_KEY"):
            print("⚠ WARNING: OPENAI_API_KEY not found in environment")
except ImportError:
    print("⚠ python-dotenv not installed. Install with: pip install python-dotenv")
    print("   Using system environment variables only")

# Add agents directory to Python path
AGENTS_DIR = Path(__file__).parent.parent / "agents"
if str(AGENTS_DIR) not in sys.path:
    sys.path.insert(0, str(AGENTS_DIR))

# Try to import agents system
# When agents directory is added to sys.path, imports work like run_multi_agent.py
try:
    from agents.team_orchestrator import TeamOrchestrator
    from agents.search_orchestrator import SearchOrchestrator
    from tools import ALL_TOOLS
    HAS_AGENTS = True
except ImportError as e:
    # Try alternative import path (direct subfolder access)
    try:
        from agents.agents.team_orchestrator import TeamOrchestrator
        from agents.agents.search_orchestrator import SearchOrchestrator
        from agents.tools import ALL_TOOLS
        HAS_AGENTS = True
    except ImportError:
        HAS_AGENTS = False
        print(f"Warning: Agents system not available: {e}")
        print("Install dependencies or check agents folder structure")
        print(f"Agents directory: {AGENTS_DIR}")

# Try to import word export libraries
try:
    from docx import Document
    from docx.shared import Pt
    from html import unescape
    import re
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

app = FastAPI(title="SidOS Backend API")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins for development
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class WordExportRequest(BaseModel):
    content: str
    file_name: Optional[str] = "Proposal.docx"


class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = "default"
    images_base64: Optional[List[str]] = None
    data_sources: Optional[Dict[str, bool]] = None


# Initialize agents system (lazy loading)
_team_orchestrator = None

def get_orchestrator():
    """Get or create the team orchestrator instance"""
    global _team_orchestrator
    if not HAS_AGENTS:
        return None
    
    if _team_orchestrator is None:
        try:
            # Create specialized agents
            search_agent = SearchOrchestrator(tools=ALL_TOOLS)
            
            # Create team orchestrator
            _team_orchestrator = TeamOrchestrator(
                specialized_agents={
                    "search": search_agent,
                }
            )
            print("✓ Agents system initialized")
        except Exception as e:
            print(f"Error initializing agents: {e}")
            return None
    
    return _team_orchestrator


def html_to_text(html_content: str) -> str:
    """Convert HTML content to plain text"""
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', html_content)
    # Decode HTML entities
    text = unescape(text)
    # Clean up extra whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()


def create_word_document(content: str, file_path: str):
    """Create a Word document from HTML content"""
    if not HAS_DOCX:
        raise HTTPException(
            status_code=500,
            detail="python-docx library not installed. Please install it: pip install python-docx"
        )
    
    doc = Document()
    
    # Convert HTML to text
    text_content = html_to_text(content)
    
    # Split into paragraphs
    paragraphs = text_content.split('\n\n')
    
    for para_text in paragraphs:
        if para_text.strip():
            # Check if it's a heading (starts with # or is all caps on first line)
            lines = para_text.split('\n')
            first_line = lines[0].strip()
            
            # Simple heading detection
            if first_line.startswith('#') or (len(first_line) < 100 and first_line.isupper()):
                # Add as heading
                heading = doc.add_heading(first_line.replace('#', '').strip(), level=1)
                # Add remaining lines as paragraphs
                for line in lines[1:]:
                    if line.strip():
                        doc.add_paragraph(line.strip())
            else:
                # Regular paragraph
                para = doc.add_paragraph(para_text.strip())
                para.style.font.size = Pt(11)
    
    # Save document
    doc.save(file_path)


@app.get("/")
def root():
    return {"message": "SidOS Backend API", "status": "running"}


@app.get("/health")
def health():
    orchestrator = get_orchestrator()
    return {
        "status": "healthy", 
        "has_docx": HAS_DOCX,
        "has_agents": HAS_AGENTS,
        "agents_enabled": orchestrator.enabled if orchestrator else False
    }


@app.post("/chat")
async def chat(request: ChatRequest):
    """
    Chat endpoint that uses the agents system
    """
    if not HAS_AGENTS:
        raise HTTPException(
            status_code=503,
            detail="Agents system not available. Check that agents folder is properly set up."
        )
    
    orchestrator = get_orchestrator()
    if not orchestrator:
        raise HTTPException(
            status_code=503,
            detail="Failed to initialize agents system"
        )
    
    if not orchestrator.enabled:
        raise HTTPException(
            status_code=503,
            detail="AI not enabled. Set OPENAI_API_KEY in environment variables."
        )
    
    start_time = time.time()
    
    try:
        # Prepare context from request
        context = {
            "session_id": request.session_id,
            "data_sources": request.data_sources or {},
        }
        
        if request.images_base64:
            context["images_base64"] = request.images_base64
        
        # Execute query through agents system
        result = orchestrator.execute(request.message, context)
        
        # Calculate latency
        latency_ms = int((time.time() - start_time) * 1000)
        
        # Extract results
        final_result = result.get("results", "")
        if isinstance(final_result, dict):
            # If results is a dict, try to get a text representation
            final_result = str(final_result)
        elif not isinstance(final_result, str):
            final_result = str(final_result)
        
        # If no results, check for error
        if not final_result and result.get("error"):
            final_result = f"Error: {result['error']}"
        
        # Format response to match frontend expectations
        response = {
            "reply": final_result,
            "session_id": request.session_id,
            "message_id": f"msg_{int(time.time() * 1000)}",
            "latency_ms": latency_ms,
            "timestamp": datetime.now().isoformat(),
        }
        
        # Add planning info if available (for debugging/logging)
        if result.get("planning"):
            response["planning_info"] = {
                "intent": result["planning"].get("intent"),
                "strategy": result["planning"].get("strategy"),
                "data_sources": result["planning"].get("data_sources", [])
            }
        
        # Add routing info if available
        if result.get("routing"):
            response["routing_info"] = {
                "agent_type": result["routing"].get("agent_type"),
                "confidence": result["routing"].get("confidence")
            }
        
        # Add thinking log from trace (for Agent Thinking panel)
        trace = result.get("trace")
        if trace and hasattr(trace, "thinking_log"):
            response["thinking_log"] = trace.thinking_log
        elif trace and isinstance(trace, dict):
            response["thinking_log"] = trace.get("thinking_log", [])
        
        return response
        
    except Exception as e:
        latency_ms = int((time.time() - start_time) * 1000)
        error_msg = f"Error processing query: {str(e)}"
        print(f"Chat error: {error_msg}")
        import traceback
        traceback.print_exc()
        
        return {
            "reply": error_msg,
            "session_id": request.session_id,
            "message_id": f"msg_{int(time.time() * 1000)}",
            "latency_ms": latency_ms,
            "timestamp": datetime.now().isoformat(),
            "error": True
        }


@app.post("/export/word")
async def export_to_word(request: WordExportRequest):
    """
    Export HTML content to Word document
    """
    try:
        # Create temporary file
        temp_dir = tempfile.gettempdir()
        file_name = request.file_name or "Proposal.docx"
        file_path = os.path.join(temp_dir, file_name)
        
        # Create Word document
        create_word_document(request.content, file_path)
        
        # Return file as download
        return FileResponse(
            file_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=file_name,
            headers={"Content-Disposition": f"attachment; filename={file_name}"}
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating Word document: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)


