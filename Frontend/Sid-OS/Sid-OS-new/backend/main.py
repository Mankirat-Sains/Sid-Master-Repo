"""
Simple FastAPI backend for SidOS Word Export
"""
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional
import os
import tempfile
from pathlib import Path
from datetime import datetime

# Try to import word export libraries
try:
    from docx import Document
    from docx.shared import Pt
    from html import unescape
    import re
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

app = FastAPI(title="SidOS Backend API")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins for development
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class WordExportRequest(BaseModel):
    content: str
    file_name: Optional[str] = "Proposal.docx"


def html_to_text(html_content: str) -> str:
    """Convert HTML content to plain text"""
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', html_content)
    # Decode HTML entities
    text = unescape(text)
    # Clean up extra whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()


def create_word_document(content: str, file_path: str):
    """Create a Word document from HTML content"""
    if not HAS_DOCX:
        raise HTTPException(
            status_code=500,
            detail="python-docx library not installed. Please install it: pip install python-docx"
        )
    
    doc = Document()
    
    # Convert HTML to text
    text_content = html_to_text(content)
    
    # Split into paragraphs
    paragraphs = text_content.split('\n\n')
    
    for para_text in paragraphs:
        if para_text.strip():
            # Check if it's a heading (starts with # or is all caps on first line)
            lines = para_text.split('\n')
            first_line = lines[0].strip()
            
            # Simple heading detection
            if first_line.startswith('#') or (len(first_line) < 100 and first_line.isupper()):
                # Add as heading
                heading = doc.add_heading(first_line.replace('#', '').strip(), level=1)
                # Add remaining lines as paragraphs
                for line in lines[1:]:
                    if line.strip():
                        doc.add_paragraph(line.strip())
            else:
                # Regular paragraph
                para = doc.add_paragraph(para_text.strip())
                para.style.font.size = Pt(11)
    
    # Save document
    doc.save(file_path)


@app.get("/")
def root():
    return {"message": "SidOS Backend API", "status": "running"}


@app.get("/health")
def health():
    return {"status": "healthy", "has_docx": HAS_DOCX}


@app.post("/export/word")
async def export_to_word(request: WordExportRequest):
    """
    Export HTML content to Word document
    """
    try:
        # Create temporary file
        temp_dir = tempfile.gettempdir()
        file_name = request.file_name or "Proposal.docx"
        file_path = os.path.join(temp_dir, file_name)
        
        # Create Word document
        create_word_document(request.content, file_path)
        
        # Return file as download
        return FileResponse(
            file_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=file_name,
            headers={"Content-Disposition": f"attachment; filename={file_name}"}
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating Word document: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)


